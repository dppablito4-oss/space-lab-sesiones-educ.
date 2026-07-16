from __future__ import annotations
import io
import re
import sys
import urllib.request
from pathlib import Path
from typing import TYPE_CHECKING, List, Optional

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

if TYPE_CHECKING:
    from main import SesionAprendizajeRequest

def set_cell_background(cell, hex_color: str):
    """Establece el color de fondo de una celda en Word."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text_direction_vertical(cell):
    """Establece la dirección del texto vertical (abajo a arriba, de izquierda a derecha) en una celda en Word."""
    tcPr = cell._tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), 'btLr')
    tcPr.append(textDirection)


def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Establece márgenes internos (padding) de una celda en dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_table_borders(table, color='CBD5E1', sz='4'):
    """Agrega bordes delgados a toda la tabla. Por defecto gris claro."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        borders.append(border)
    tblPr.append(borders)


def set_table_col_widths(table, widths_twip: list):
    """
    Fuerza los anchos de columna a nivel XML (w:tblGrid + w:gridCol)
    para que Word respete exactamente los anchos indicados en twips (1 inch = 1440 twips).
    """
    tbl = table._tbl
    # Eliminar tblGrid existente si hay
    for old_grid in tbl.findall(qn('w:tblGrid')):
        tbl.remove(old_grid)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_twip:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w)))
        tblGrid.append(gridCol)
    # Insertar tblGrid después de tblPr
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)
    # Actualizar w:tcW de cada celda en la fila 0 como referencia
    for row in tbl.findall(qn('w:tr')):
        tcs = row.findall(qn('w:tc'))
        for idx, tc in enumerate(tcs):
            if idx < len(widths_twip):
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(int(widths_twip[idx])))
                tcW.set(qn('w:type'), 'dxa')


def add_table_borders_black(table):
    """Agrega bordes negros medios a toda la tabla (estilo plantilla oficial)."""
    add_table_borders(table, color='000000', sz='8')


def set_run_color(run, hex_color: str):
    """Aplica color hex a un run de texto Word."""
    run.font.color.rgb = RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16)
    )


def set_cell_text_white_bold(cell, text: str, font_size_pt: float = 9):
    """Escribe texto en blanco y negrita en la primera línea de una celda."""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Limpiar párrafo
    for run in p.runs:
        run.text = ''
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def build_docx_from_html(html_content: str) -> io.BytesIO:
    """Parsea recursivamente la estructura HTML y construye un documento .docx nativo."""
    doc = Document()
    
    # Configuración de márgenes estándar
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(30, 41, 59) # Slate-800

    soup = BeautifulSoup(html_content, 'html.parser')

    # Eliminar katex-html para evitar texto duplicado
    for katex_html in soup.find_all(class_='katex-html'):
        katex_html.decompose()

    # Limpiar elementos interactivos del editor que no pertenecen al documento
    for selector in ['no-print', 'add-logo-placeholder', 'btn-remove-logo']:
        for el in soup.find_all(class_=selector):
            el.decompose()

    processed_tags = set()

    def add_runs_to_paragraph(paragraph, element, is_bold=False, is_italic=False):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    if is_bold:
                        run.bold = True
                    if is_italic:
                        run.italic = True
            elif isinstance(child, Tag):
                if child.name == 'strong' or child.name == 'b':
                    add_runs_to_paragraph(paragraph, child, is_bold=True, is_italic=is_italic)
                elif child.name == 'em' or child.name == 'i':
                    add_runs_to_paragraph(paragraph, child, is_bold=is_bold, is_italic=True)
                elif child.name == 'span':
                    # Manejar estilos locales en spans si fuera necesario
                    add_runs_to_paragraph(paragraph, child, is_bold=is_bold, is_italic=is_italic)
                elif child.name == 'br':
                    paragraph.add_run('\n')
                else:
                    add_runs_to_paragraph(paragraph, child, is_bold=is_bold, is_italic=is_italic)

    def walk_tree(element):
        if element in processed_tags:
            return
        
        if isinstance(element, Tag):
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                
                sz = 14 if element.name == 'h1' else 12
                run = p.add_run(element.get_text().strip())
                run.bold = True
                run.font.size = Pt(sz)
                processed_tags.add(element)
                
            elif element.name == 'p':
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                add_runs_to_paragraph(p, element)
                processed_tags.add(element)
                
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.15
                    add_runs_to_paragraph(p, li)
                processed_tags.add(element)
                
            elif element.name == 'table':
                # Reconstruir la tabla HTML de manera exacta
                rows = element.find_all('tr')
                if not rows:
                    processed_tags.add(element)
                    return
                
                # Encontrar dimensiones
                max_cols = 0
                for row in rows:
                    cells = row.find_all(['td', 'th'])
                    max_cols = max(max_cols, len(cells))
                
                if max_cols == 0:
                    processed_tags.add(element)
                    return
                
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.autofit = True
                add_table_borders(table)
                
                for r_idx, row in enumerate(rows):
                    html_cells = row.find_all(['td', 'th'])
                    for c_idx, html_cell in enumerate(html_cells):
                        if c_idx >= max_cols:
                            continue
                        
                        cell = table.cell(r_idx, c_idx)
                        is_header = html_cell.name == 'th' or 'th' in html_cell.get('class', [])
                        
                        bg_color = "FFFFFF"
                        style_attr = html_cell.get('style', '')
                        class_attr = html_cell.get('class', [])

                        hex_match = re.search(r'background-color:\s*#([A-Fa-f0-9]{6})', style_attr)
                        if hex_match:
                            bg_color = hex_match.group(1).upper()
                        elif 'cell-peru' in class_attr:
                            bg_color = "C0392B"
                        elif 'cell-minedu' in class_attr:
                            bg_color = "2C3E50"

                        set_cell_background(cell, bg_color)
                        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)

                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.1

                        if is_header:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.bold = True
                            if bg_color in ["C0392B", "2C3E50"]:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                run.font.color.rgb = RGBColor(15, 23, 42)
                            add_runs_to_paragraph(p, html_cell)
                        else:
                            add_runs_to_paragraph(p, html_cell)
                
                processed_tags.add(element)
                
            else:
                for child in element.children:
                    walk_tree(child)
        else:
            # Es un string navegable suelto
            pass

    for child in soup.body.children if soup.body else soup.children:
        walk_tree(child)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def get_image_stream(url: str):
    """Intenta descargar la imagen desde la URL y devuelve un BytesIO. Retorna None si falla o es vacía."""
    if not url or not (url.startswith("http://") or url.startswith("https://")):
        return None
    try:
        req = urllib.request.Request(
            url, 
            headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        )
        with urllib.request.urlopen(req, timeout=3) as response:
            return io.BytesIO(response.read())
    except Exception as e:
        print(f"[WARN LOGO] No se pudo descargar el logo {url}: {e}")
        return None


def format_latex_to_unicode(text: str) -> str:
    """
    Traduce expresiones comunes de LaTeX a caracteres Unicode matemáticos para legibilidad en Word.
    """
    if not text:
        return ""
    
    # Tabla de equivalencias directas de LaTeX a Unicode
    # ── Convertir tablas LaTeX \begin{array}...\end{array} a texto tabular ──
    def convert_latex_table(match):
        content = match.group(1)
        # Obtener filas
        rows_raw = re.split(r'\\\\', content)
        rows = []
        for row_raw in rows_raw:
            row_raw = row_raw.strip()
            if not row_raw:
                continue
            # Separar celdas por &
            cells_raw = row_raw.split('&')
            cells = [c.strip() for c in cells_raw]
            rows.append(cells)
        if not rows:
            return text
        num_cols = max(len(r) for r in rows)
        # Calcular anchos de columna
        col_widths = []
        for c in range(num_cols):
            w = max((len(r[c]) if c < len(r) else 0) for r in rows)
            col_widths.append(max(w, 4))
        # Construir tabla unicode
        lines = []
        sep = '─' * (sum(col_widths) + num_cols * 3 + 1)
        lines.append(sep)
        for r_idx, row in enumerate(rows):
            line_parts = []
            for c_idx in range(num_cols):
                cell_val = row[c_idx] if c_idx < len(row) else ''
                line_parts.append(cell_val.ljust(col_widths[c_idx]))
            lines.append('│ ' + ' │ '.join(line_parts) + ' │')
            if r_idx == 0:
                lines.append(sep)
        lines.append(sep)
        return '\n'.join(lines)

    text = re.sub(r'\\begin\{(?:array|tabular)\}(?:\{[^}]*\})?(.+?)\\end\{(?:array|tabular)\}', convert_latex_table, text, flags=re.DOTALL)

    replacements = {
        r'\pm': '±',
        r'\le': '≤',
        r'\ge': '≥',
        r'\neq': '≠',
        r'\times': '×',
        r'\div': '÷',
        r'\approx': '≈',
        r'\alpha': 'α',
        r'\beta': 'β',
        r'\gamma': 'γ',
        r'\theta': 'θ',
        r'\pi': 'π',
        r'\infty': '∞',
        r'\Delta': 'Δ',
        r'\Sigma': 'Σ',
        r'\rightarrow': '→',
        r'\leftarow': '←',
        r'\leftrightarrow': '↔',
        r'\Rightarrow': '⇒',
        r'\Leftarrow': '⇐',
        r'\partial': '∂',
        r'\sqrt': '√',
        r'\cdot': '·',
        r'\text{o}': ' o ',
        r'\;': ' ',
        r'\:': ' ',
        r'\,': ' ',
        r'\!': '',
    }
    
    # Remover delimitadores de dólar
    cleaned = text
    cleaned = re.sub(r'\$\$(.*?)\$\$', r'\1', cleaned)
    cleaned = re.sub(r'\$(.*?)\$', r'\1', cleaned)
    
    # Aplicar equivalencias de comandos
    for key, val in replacements.items():
        cleaned = cleaned.replace(key, val)
        
    # Reemplazos de potencias comunes a superíndices unicode
    supers = {'0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴', '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹', 
              '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾', 'n': 'ⁿ', 'x': 'ˣ', 'y': 'ʸ', 'i': 'ⁱ'}
    
    def replace_super(match):
        val = match.group(1)
        return "".join(supers.get(c, c) for c in val)
        
    cleaned = re.sub(r'\^\{?([0-9+\-xyn]+)\}?', replace_super, cleaned)
    
    # Reemplazos de subíndices comunes a unicode
    subs = {'0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄', '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉', 
            '+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎', 'n': 'ₙ', 'x': 'ₓ', 'y': 'ᵧ', 'i': 'ᵢ'}
            
    def replace_sub(match):
        val = match.group(1)
        return "".join(subs.get(c, c) for c in val)
        
    cleaned = re.sub(r'\_\{?([0-9+\-xyni]+)\}?', replace_sub, cleaned)
    
    return cleaned


def preprocess_session_latex(session):
    """
    Recorre el modelo de la sesión y traduce todas las expresiones LaTeX de sus textos
    a caracteres Unicode matemáticos listos para la exportación a Word.
    """
    if not session:
        return
        
    # Metadatos
    if session.metadata:
        session.metadata.titulo = format_latex_to_unicode(session.metadata.titulo)
        session.metadata.institucion = format_latex_to_unicode(session.metadata.institucion)
        session.metadata.docente = format_latex_to_unicode(session.metadata.docente)
        session.metadata.area = format_latex_to_unicode(session.metadata.area)
        session.metadata.unidad = format_latex_to_unicode(session.metadata.unidad)
        
    # Propósitos
    if session.proposito:
        session.proposito.proposito_texto = format_latex_to_unicode(session.proposito.proposito_texto)
        session.proposito.conocimientos = format_latex_to_unicode(session.proposito.conocimientos)
        session.proposito.competencia = format_latex_to_unicode(session.proposito.competencia)
        session.proposito.estandar = format_latex_to_unicode(session.proposito.estandar)
        session.proposito.producto_evidencia = format_latex_to_unicode(session.proposito.producto_evidencia)
        session.proposito.instrumento = format_latex_to_unicode(session.proposito.instrumento)
        
        if session.proposito.capacidades:
            session.proposito.capacidades = [format_latex_to_unicode(c) for c in session.proposito.capacidades]
        if session.proposito.criterios:
            session.proposito.criterios = [format_latex_to_unicode(cr) for cr in session.proposito.criterios]

    # Competencias transversales
    if session.competencias_transversales:
        for ct in session.competencias_transversales:
            ct.titulo = format_latex_to_unicode(ct.titulo)
            if ct.desempenos:
                ct.desempenos = [format_latex_to_unicode(d) for d in ct.desempenos]

    # Enfoques transversales
    if session.enfoques_transversales:
        for et in session.enfoques_transversales:
            et.nombre = format_latex_to_unicode(et.nombre)
            et.valor = format_latex_to_unicode(et.valor)
            et.actitudes = format_latex_to_unicode(et.actitudes)

    # Recursos
    if session.recursos:
        session.recursos.enlaces = format_latex_to_unicode(session.recursos.enlaces)
        session.recursos.materiales = format_latex_to_unicode(session.recursos.materiales)
        session.recursos.refuerzo = format_latex_to_unicode(session.recursos.refuerzo)

    # Momentos
    if getattr(session, 'momentos', None):
        momentos = session.momentos
        # Inicio
        if getattr(momentos, 'inicio', None) and getattr(momentos.inicio, 'actividades', None):
            momentos.inicio.actividades = [format_latex_to_unicode(a) for a in momentos.inicio.actividades]
        # Desarrollo
        if getattr(momentos, 'desarrollo', None) and getattr(momentos.desarrollo, 'procesos', None):
            for proc in momentos.desarrollo.procesos:
                if getattr(proc, 'titulo', None):
                    proc.titulo = format_latex_to_unicode(proc.titulo)
                if getattr(proc, 'contenido', None):
                    proc.contenido = [format_latex_to_unicode(c) for c in proc.contenido]
        # Cierre
        if getattr(momentos, 'cierre', None):
            cierre = momentos.cierre
            if getattr(cierre, 'metacognicion', None):
                cierre.metacognicion = [format_latex_to_unicode(m) for m in cierre.metacognicion]
            if getattr(cierre, 'evaluacion', None):
                cierre.evaluacion = [format_latex_to_unicode(ev) for ev in cierre.evaluacion]
            if getattr(cierre, 'extension', None):
                cierre.extension = [format_latex_to_unicode(ex) for ex in cierre.extension]

    # Ficha de trabajo
    if getattr(session, 'ficha_trabajo', None):
        ficha = session.ficha_trabajo
        if getattr(ficha, 'titulo', None):
            ficha.titulo = format_latex_to_unicode(ficha.titulo)
        if getattr(ficha, 'indicaciones', None):
            ficha.indicaciones = format_latex_to_unicode(ficha.indicaciones)
        if getattr(ficha, 'actividades', None):
            ficha.actividades = format_latex_to_unicode(ficha.actividades)

    # Firmas (si existieran dinámicamente)
    if getattr(session, 'firmas', None):
        for f in session.firmas:
            if getattr(f, 'nombre', None):
                f.nombre = format_latex_to_unicode(f.nombre)
            if getattr(f, 'cargo', None):
                f.cargo = format_latex_to_unicode(f.cargo)


def build_docx_from_json(session: SesionAprendizajeRequest) -> io.BytesIO:
    # Preprocesar LaTeX a Unicode para asegurar compatibilidad matemática y estética en Word
    preprocess_session_latex(session)
    
    doc = Document()
    
    # Margenes exactos A4 y A4 page size
    for s in doc.sections:
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = Inches(0.1)
        s.bottom_margin = Inches(0.1)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # Estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(30, 41, 59)

    # Colores exactos de la plantilla XML
    PEACH = 'F7CAAC'
    BLUE_HDR = 'BDD6EE'
    YELLOW_HDR = 'FFE599'
    GRAY_VAL = 'F2F2F2'
    PEACH_MOM = 'FBE5D5'
    GRAY_MOM = 'EDEDED'
    YELLOW_VAL = 'FFF2CC'
    BULLET_COLORS = ['2980B9', 'C0392B', '27AE60', '8E44AD', '16A085']

    def _label(cell, text):
        set_cell_background(cell, PEACH)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)

    def _val(cell, text):
        set_cell_background(cell, GRAY_VAL)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.1
        p.add_run(text or "").font.size = Pt(9)

    def _hdr(cell, text, bg=BLUE_HDR, sz=9):
        set_cell_background(cell, bg)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(sz)

    def _bullet_cell(cell, items):
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        cell.text = ""
        for i, item in enumerate(items):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            bc = BULLET_COLORS[i % len(BULLET_COLORS)]
            rb = p.add_run(u"\u25cf ")
            rb.font.size = Pt(9)
            rb.font.color.rgb = RGBColor(int(bc[0:2], 16), int(bc[2:4], 16), int(bc[4:6], 16))
            p.add_run(item).font.size = Pt(9)

    def make_vertical_text(text: str) -> str:
        # Stack characters vertically separated by double newlines to match template
        return "\n" + "\n\n".join(list(text)) + "\n"

    def _write_momento_cell(cell, nombre, sub_bullets, tiempo):
        set_cell_background(cell, GRAY_MOM)
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p_main = cell.paragraphs[0]
        rm = p_main.add_run(nombre)
        rm.bold = True
        rm.font.size = Pt(10)
        if sub_bullets:
            p_sub = cell.add_paragraph()
            p_sub.paragraph_format.space_before = Pt(4)
            rs = p_sub.add_run(sub_bullets)
            rs.font.size = Pt(8)
            rs.font.color.rgb = RGBColor(71, 85, 105)
        if tiempo:
            p_t = cell.add_paragraph()
            p_t.paragraph_format.space_before = Pt(6)
            rt = p_t.add_run(u"\u23f1 TIEMPO: " + tiempo + " min")
            rt.bold = True
            rt.font.size = Pt(8.5)
            rt.font.color.rgb = RGBColor(192, 57, 43)

    def _write_vertical_cell(cell, txt):
        set_cell_background(cell, PEACH_MOM)
        set_cell_margins(cell, top=120, bottom=120, left=40, right=40)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(make_vertical_text(txt))
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(120, 60, 20) # Cafe oscuro / oxidado

    # ===============================================================
    # CABECERA INSTITUCIONAL EN EL HEADER NATIVO (TODAS LAS PÁGINAS)
    # ===============================================================
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header = section.header
    header.is_linked_to_previous = False

    # Cargar logo genérico local de la marca (assets/logo.png) para ambos lados
    logo_left_stream = None
    logo_right_stream = None
    try:
        logo_path = Path(__file__).resolve().parent.parent / "assets" / "logo.png"
        if logo_path.exists():
            with open(logo_path, "rb") as f:
                logo_bytes = f.read()
                logo_left_stream = io.BytesIO(logo_bytes)
                logo_right_stream = io.BytesIO(logo_bytes)
    except Exception as e:
        print(f"[WARN] No se pudo cargar logo genérico local: {e}")

    # Preparar textos institucionales dinámicos
    dre_txt = (session.metadata.dre or "DRE").strip()
    ugel_txt = (session.metadata.ugel or "UGEL").strip()
    ie_txt = (session.metadata.institucion or "").strip()
    # Línea institucional central: «IE - DRE - UGEL»
    partes_ie = [p for p in [ie_txt, dre_txt, ugel_txt] if p]
    linea_ie = " - ".join(partes_ie) if partes_ie else "Institución Educativa"

    # Un solo párrafo centrado con el logo de la marca (Space Lab)
    p_logo = header.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(4)

    if logo_left_stream:
        try:
            run_logo = p_logo.add_run()
            run_logo.add_picture(logo_left_stream, height=Inches(0.55))
        except Exception:
            run_lbl = p_logo.add_run("🚀 Space Lab")
            run_lbl.bold = True
            run_lbl.font.size = Pt(11)
            run_lbl.font.color.rgb = RGBColor(56, 189, 248)
    else:
        run_lbl = p_logo.add_run("🚀 Space Lab")
        run_lbl.bold = True
        run_lbl.font.size = Pt(11)
        run_lbl.font.color.rgb = RGBColor(56, 189, 248)

    # Línea divisoria inferior de 1.5 pt
    p_div = header.add_paragraph()
    p_div.paragraph_format.space_before = Pt(3)
    p_div.paragraph_format.space_after = Pt(6)
    pBrd = OxmlElement('w:pBrd')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '12')  # 1.5 pt
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '334155')  # Slate 700
    pBrd.append(bottom_border)
    p_div._p.get_or_add_pPr().append(pBrd)

    # ===============================================================
    # TITULO DE LA SESIÓN DE APRENDIZAJE
    # ===============================================================
    t_tbl = doc.add_table(rows=1, cols=1)
    t_tbl.autofit = False
    t_tbl.rows[0].cells[0].width = Inches(6.77)
    add_table_borders_black(t_tbl)
    
    title_text = f"SESIÓN DE APRENDIZAJE N° {session.metadata.numero_sesion or '01'}"
    set_cell_text_white_bold(t_tbl.cell(0, 0), title_text, font_size_pt=11.5)
    set_cell_background(t_tbl.cell(0, 0), '2980B9') # Azul Principal
    set_cell_margins(t_tbl.cell(0, 0), top=120, bottom=120, left=180, right=180)

    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ===============================================================
    # TABLA DATOS INFORMATIVOS — Diseño limpio 4 columnas (label | val | label | val)
    # ===============================================================
    # Anchos en twips: col_lbl=1584 (~1.1in), col_val=2376 (~1.65in) × 2 pares = 7920 twips total (~5.5in)
    # Nota: 6.77in = 9748 twips; ajustamos para que quepa todo el ancho útil.
    _DI_LBL = 1584   # ~1.1 in
    _DI_VAL = 3288   # ~2.28 in
    _DI_TWIPS = [_DI_LBL, _DI_VAL, _DI_LBL, _DI_VAL]  # total ~9744 twips ≈ 6.77in

    di = doc.add_table(rows=4, cols=4)
    di.autofit = False
    add_table_borders_black(di)
    set_table_col_widths(di, _DI_TWIPS)

    def _di_label(cell, text):
        """Celda de etiqueta PEACH con texto bold."""
        set_cell_background(cell, PEACH)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(30, 41, 59)

    def _di_val(cell, text, bold=False, sz=9.5):
        """Celda de valor gris claro."""
        set_cell_background(cell, GRAY_VAL)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(text or "")
        r.bold = bold
        r.font.size = Pt(sz)

    # Fila 0: Institución Educativa  |  (valor)  |  Nivel  |  (valor)
    _di_label(di.cell(0, 0), "Institución Educativa")
    _di_val(di.cell(0, 1), session.metadata.institucion or "No especificada")
    _di_label(di.cell(0, 2), "Nivel")
    _di_val(di.cell(0, 3), session.metadata.nivel or "SECUNDARIA")

    # Fila 1: Docente  |  (valor)  |  Área  |  (valor)
    _di_label(di.cell(1, 0), "Docente")
    _di_val(di.cell(1, 1), session.metadata.docente or "No especificado")
    _di_label(di.cell(1, 2), "Área")
    _di_val(di.cell(1, 3), session.metadata.area or "No especificada")

    # Fila 2: Grado  |  (valor)  |  Unidad/Proyecto  |  (valor)
    _di_label(di.cell(2, 0), "Grado y Sección")
    grado_seccion = " ".join(filter(None, [
        session.metadata.grado or "",
        session.metadata.seccion or ""
    ])) or "No especificado"
    _di_val(di.cell(2, 1), grado_seccion)
    _di_label(di.cell(2, 2), "Unidad / Proyecto")
    _di_val(di.cell(2, 3), session.metadata.unidad or "No especificada")

    # Fila 3: Fecha  |  (valor)  |  Duración  |  (valor)
    _di_label(di.cell(3, 0), "Fecha")
    _di_val(di.cell(3, 1), session.metadata.fecha or "No especificada")
    _di_label(di.cell(3, 2), "Duración (min)")
    _di_val(di.cell(3, 3), session.metadata.duracion or "90 min")

    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ===============================================================
    # TABLA TÍTULO DE SESIÓN, PROPÓSITO Y CONOCIMIENTOS
    # ===============================================================
    pc = doc.add_table(rows=6, cols=1)
    pc.autofit = False
    pc.rows[0].cells[0].width = Inches(6.77)
    add_table_borders_black(pc)

    _hdr(pc.cell(0, 0), "TÍTULO DE LA SESIÓN", bg=BLUE_HDR, sz=9)
    
    cell_tit = pc.cell(1, 0)
    set_cell_margins(cell_tit, top=80, bottom=80, left=120, right=120)
    p_tit = cell_tit.paragraphs[0]
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tit = p_tit.add_run(session.metadata.titulo or "Título de la sesión de aprendizaje")
    r_tit.bold = True
    r_tit.font.size = Pt(10.5)
    r_tit.font.color.rgb = RGBColor(192, 57, 43)

    _hdr(pc.cell(2, 0), "PROPÓSITO DE LA SESIÓN:", bg=BLUE_HDR, sz=9)
    
    cell_prop = pc.cell(3, 0)
    set_cell_margins(cell_prop, top=100, bottom=100, left=120, right=120)
    p_prop = cell_prop.paragraphs[0]
    p_prop.paragraph_format.line_spacing = 1.15
    p_prop.add_run(session.proposito.proposito_texto or "No especificado").font.size = Pt(9.5)

    _hdr(pc.cell(4, 0), "CONOCIMIENTOS:", bg=PEACH, sz=9)
    
    cell_con = pc.cell(5, 0)
    set_cell_margins(cell_con, top=100, bottom=100, left=120, right=120)
    p_con = cell_con.paragraphs[0]
    p_con.paragraph_format.line_spacing = 1.15
    p_con.add_run(session.proposito.conocimientos or "No especificado").font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ===============================================================
    # TABLA PROPÓSITOS DE APRENDIZAJE (COMPETENCIAS, CAPACIDADES, CRITERIOS)
    # ===============================================================
    pa = doc.add_table(rows=6, cols=5)
    pa.autofit = False
    add_table_borders_black(pa)

    c_hdr = pa.cell(0, 0).merge(pa.cell(0, 4))
    _hdr(c_hdr, "PROPÓSITOS DE APRENDIZAJE", bg=BLUE_HDR, sz=10)

    c_comp = pa.cell(1, 0).merge(pa.cell(1, 4))
    set_cell_background(c_comp, 'FFFFFF')
    set_cell_margins(c_comp, top=80, bottom=80, left=140, right=140)
    p_comp = c_comp.paragraphs[0]
    p_comp.paragraph_format.line_spacing = 1.15
    rc1 = p_comp.add_run("Competencia: ")
    rc1.bold = True
    rc1.font.size = Pt(9.5)
    rc2 = p_comp.add_run(session.proposito.competencia or "No especificada")
    rc2.font.size = Pt(9.5)

    c_est = pa.cell(2, 0).merge(pa.cell(2, 4))
    set_cell_background(c_est, 'FFFFFF')
    set_cell_margins(c_est, top=80, bottom=80, left=140, right=140)
    p_est = c_est.paragraphs[0]
    p_est.paragraph_format.line_spacing = 1.15
    re1 = p_est.add_run("Estándar de aprendizaje: ")
    re1.bold = True
    re1.font.size = Pt(9.5)
    re2 = p_est.add_run(session.proposito.estandar or "No especificado")
    re2.font.size = Pt(9.5)

    headers_pa = ["COMPETENCIAS", "CAPACIDADES", "CRITERIOS DE EVALUACIÓN", "PRODUCTO / EVIDENCIA", "INSTRUMENTOS DE EVALUACIÓN"]
    for i, ht in enumerate(headers_pa):
        _hdr(pa.cell(3, i), ht, bg=BLUE_HDR, sz=8)

    c_comp_v = pa.cell(4, 0).merge(pa.cell(5, 0))
    set_cell_background(c_comp_v, 'FFFFFF')
    set_cell_margins(c_comp_v, top=100, bottom=100, left=100, right=100)
    p_cv = c_comp_v.paragraphs[0]
    p_cv.paragraph_format.line_spacing = 1.1
    p_cv.add_run(session.proposito.competencia or "No especificada").font.size = Pt(8.5)

    _bullet_cell(pa.cell(4, 1), session.proposito.capacidades)
    _bullet_cell(pa.cell(4, 2), session.proposito.criterios)

    c_ev = pa.cell(4, 3)
    set_cell_margins(c_ev, top=100, bottom=100, left=100, right=100)
    p_ev = c_ev.paragraphs[0]
    p_ev.paragraph_format.line_spacing = 1.1
    p_ev.add_run(session.proposito.producto_evidencia or "No especificado").font.size = Pt(8.5)

    c_ins_v = pa.cell(4, 4).merge(pa.cell(5, 4))
    set_cell_background(c_ins_v, 'F2F2F2')
    set_cell_margins(c_ins_v, top=100, bottom=100, left=100, right=100)
    p_iv = c_ins_v.paragraphs[0]
    p_iv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_iv.paragraph_format.line_spacing = 1.1
    ri_v = p_iv.add_run(session.proposito.instrumento or "Lista de Cotejo")
    ri_v.bold = True
    ri_v.font.size = Pt(8.5)

    set_cell_margins(pa.cell(5, 1), top=40, bottom=40, left=40, right=40)
    set_cell_margins(pa.cell(5, 2), top=40, bottom=40, left=40, right=40)
    set_cell_margins(pa.cell(5, 3), top=40, bottom=40, left=40, right=40)

    anchos_pa = [Inches(1.2), Inches(1.37), Inches(2.2), Inches(1.2), Inches(0.8)]
    for row in pa.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = anchos_pa[ci]

    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ===============================================================
    # TABLA COMPETENCIAS TRANSVERSALES
    # ===============================================================
    cts = session.competencias_transversales
    n_ct = len(cts) if cts else 2
    ct_tbl = doc.add_table(rows=1 + n_ct, cols=2)
    ct_tbl.autofit = False
    add_table_borders_black(ct_tbl)

    _hdr(ct_tbl.cell(0, 0), "COMPETENCIAS TRANSVERSALES", bg=BLUE_HDR, sz=9)
    _hdr(ct_tbl.cell(0, 1), "DESEMPENOS PRECISADOS / PRODUCTO / INSTRUMENTOS", bg=BLUE_HDR, sz=9)

    if cts:
        for ci, ct in enumerate(cts):
            rn = 1 + ci
            set_cell_background(ct_tbl.cell(rn, 0), 'F2F2F2')
            set_cell_margins(ct_tbl.cell(rn, 0), top=100, bottom=100, left=120, right=120)
            ct_tbl.cell(rn, 0).paragraphs[0].add_run(ct.titulo or "Competencia Transversal").font.size = Pt(8.5)
            _bullet_cell(ct_tbl.cell(rn, 1), ct.desempenos if ct.desempenos else ["No especificado"])
    else:
        defaults_ct = [
            ("Gestiona su aprendizaje de manera autonoma", [
                "Define metas de aprendizaje para alcanzar sus objetivos pedagógicos.",
                "Organiza acciones estrategicas para alcanzar sus metas de aprendizaje.",
                "Monitorea y ajusta su desempeno durante el proceso de aprendizaje."
            ]),
            ("Se desenvuelve en los entornos virtuales generados por las TIC", [
                "Personaliza entornos virtuales segun sus necesidades de indagacion.",
                "Gestiona informacion del entorno virtual de manera segura.",
                "Interactua en entornos virtuales y crea objetos virtuales."
            ])
        ]
        for ci, (title, items) in enumerate(defaults_ct):
            rn = 1 + ci
            set_cell_background(ct_tbl.cell(rn, 0), 'F2F2F2')
            set_cell_margins(ct_tbl.cell(rn, 0), top=100, bottom=100, left=120, right=120)
            ct_tbl.cell(rn, 0).paragraphs[0].add_run(title).font.size = Pt(8.5)
            _bullet_cell(ct_tbl.cell(rn, 1), items)

    anchos_ct = [Inches(2.5), Inches(4.27)]
    for row in ct_tbl.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = anchos_ct[ci]

    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ===============================================================
    # TABLA ENFOQUES TRANSVERSALES
    # ===============================================================
    enfoques = session.enfoques_transversales
    n_enf = len(enfoques) if enfoques else 2
    et_tbl = doc.add_table(rows=1 + n_enf, cols=3)
    et_tbl.autofit = False
    add_table_borders_black(et_tbl)

    headers_et = ["Enfoque(s) transversal(es)", "Valores", "Actitudes o acciones observables"]
    for i, ht in enumerate(headers_et):
        _hdr(et_tbl.cell(0, i), ht, bg=BLUE_HDR, sz=9)

    if enfoques:
        for ci, et in enumerate(enfoques):
            rn = 1 + ci
            set_cell_background(et_tbl.cell(rn, 0), 'F2F2F2')
            set_cell_margins(et_tbl.cell(rn, 0), top=80, bottom=80, left=120, right=120)
            et_tbl.cell(rn, 0).paragraphs[0].add_run(et.nombre or "Enfoque Transversal").font.size = Pt(8.5)
            
            set_cell_background(et_tbl.cell(rn, 1), 'FFF2CC')
            set_cell_margins(et_tbl.cell(rn, 1), top=80, bottom=80, left=120, right=120)
            et_tbl.cell(rn, 1).paragraphs[0].add_run(et.valor or "No especificado").font.size = Pt(8.5)
            
            set_cell_background(et_tbl.cell(rn, 2), 'FFF2CC')
            set_cell_margins(et_tbl.cell(rn, 2), top=80, bottom=80, left=120, right=120)
            et_tbl.cell(rn, 2).paragraphs[0].add_run(et.actitudes or "No especificadas").font.size = Pt(8.5)
    else:
        defaults_et = [
            ("Enfoque Ambiental", "Justicia y solidaridad", "Reduce el uso de materiales desechables, reutilizando cuadernos, hojas y envases cuando sea posible durante las actividades del aula."),
            ("Enfoque Busqueda de la Excelencia", "Equidad y Justicia", "Dialoga con tus compañeros para resolver desacuerdos y escucha con atencion.")
        ]
        for ci, (name, val, act) in enumerate(defaults_et):
            rn = 1 + ci
            set_cell_background(et_tbl.cell(rn, 0), 'F2F2F2')
            set_cell_margins(et_tbl.cell(rn, 0), top=80, bottom=80, left=120, right=120)
            et_tbl.cell(rn, 0).paragraphs[0].add_run(name).font.size = Pt(8.5)
            
            set_cell_background(et_tbl.cell(rn, 1), 'FFF2CC')
            set_cell_margins(et_tbl.cell(rn, 1), top=80, bottom=80, left=120, right=120)
            et_tbl.cell(rn, 1).paragraphs[0].add_run(val).font.size = Pt(8.5)
            
            set_cell_background(et_tbl.cell(rn, 2), 'FFF2CC')
            set_cell_margins(et_tbl.cell(rn, 2), top=80, bottom=80, left=120, right=120)
            et_tbl.cell(rn, 2).paragraphs[0].add_run(act).font.size = Pt(8.5)

    anchos_et = [Inches(2.0), Inches(1.37), Inches(3.4)]
    for row in et_tbl.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = anchos_et[ci]

    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ===============================================================
    # TABLA RECURSOS Y MATERIALES
    # ===============================================================
    if getattr(session, 'recursos', None):
        rec = doc.add_table(rows=4, cols=2)
        rec.autofit = False
        add_table_borders_black(rec)

        _hdr(rec.cell(0, 0), "Páginas de: Texto de, otros textos de consulta/ Enlace web, etc.", bg=YELLOW_HDR, sz=8.5)
        set_cell_background(rec.cell(0, 1), 'F2F2F2')
        set_cell_margins(rec.cell(0, 1), top=80, bottom=80, left=120, right=120)
        rec.cell(0, 1).paragraphs[0].add_run(session.recursos.enlaces or "https://www.perueduca.pe/#/home/materiales-educativos").font.size = Pt(8.5)

        _hdr(rec.cell(1, 0), "", bg=YELLOW_HDR, sz=8.5)
        set_cell_background(rec.cell(1, 1), 'F2F2F2')
        set_cell_margins(rec.cell(1, 1), top=40, bottom=40, left=120, right=120)

        _hdr(rec.cell(2, 0), "Materiales y recursos", bg=PEACH_MOM, sz=8.5)
        set_cell_background(rec.cell(2, 1), 'F2F2F2')
        set_cell_margins(rec.cell(2, 1), top=80, bottom=80, left=120, right=120)
        rec.cell(2, 1).paragraphs[0].add_run(session.recursos.materiales or "Ficha de actividades N° 01-02").font.size = Pt(8.5)

        _hdr(rec.cell(3, 0), "Actividades de Refuerzo Escolar (N° ficha y Título)", bg=YELLOW_HDR, sz=8.5)
        set_cell_background(rec.cell(3, 1), 'F2F2F2')
        set_cell_margins(rec.cell(3, 1), top=80, bottom=80, left=120, right=120)
        rec.cell(3, 1).paragraphs[0].add_run(session.recursos.refuerzo or "No especificado").font.size = Pt(8.5)

        anchos_rec = [Inches(2.5), Inches(4.27)]
        for row in rec.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = anchos_rec[ci]

        doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ===============================================================
    # TABLA SECUENCIAL DE MOMENTOS DE LA SESIÓN (DISEÑO 4 COLUMNAS PREMIUM)
    # ===============================================================
    if getattr(session, 'momentos', None):
        procs = session.momentos.desarrollo.procesos if getattr(session.momentos, 'desarrollo', None) and getattr(session.momentos.desarrollo, 'procesos', None) else []
        n_proc = len(procs) if procs else 1
        n_rows = 1 + 1 + n_proc + 1

        mt = doc.add_table(rows=n_rows, cols=4)
        mt.autofit = False
        add_table_borders_black(mt)

        _hdr(mt.cell(0, 0), "MOMENTOS DE LA SESION", bg=BLUE_HDR, sz=9.5)

        cell_est_hdr = mt.cell(0, 1).merge(mt.cell(0, 3))
        _hdr(cell_est_hdr, "ESTRATEGIAS / ACTIVIDADES", bg=BLUE_HDR, sz=9.5)

        # Fila 1: INICIO
        _write_momento_cell(mt.cell(1, 0), "INICIO",
            "Saberes Previos\nProblematizacion\nMotivacion",
            session.momentos.inicio.tiempo_total or "")
        _write_vertical_cell(mt.cell(1, 1), "MOTIVACION")
        _write_vertical_cell(mt.cell(1, 3), "EVALUACION")

        cell_ini = mt.cell(1, 2)
        set_cell_margins(cell_ini, top=120, bottom=120, left=140, right=140)
        cell_ini.text = ""
        for act in (session.momentos.inicio.actividades or []):
            p = cell_ini.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            p.add_run(act).font.size = Pt(9.5)

        # Filas 2 a 2+n_proc-1: DESARROLLO
        _write_momento_cell(mt.cell(2, 0), "DESARROLLO",
            "Gestion y Acompanamiento del Desarrollo de las Competencias\n(Procesos didacticos del Area)",
            session.momentos.desarrollo.tiempo_total or "")
        _write_vertical_cell(mt.cell(2, 1), "MOTIVACION")
        _write_vertical_cell(mt.cell(2, 3), "EVALUACION")

        for idx in range(n_proc):
            rn = 2 + idx
            cell_des = mt.cell(rn, 2)
            set_cell_margins(cell_des, top=120, bottom=120, left=140, right=140)
            cell_des.text = ""
            if procs:
                proc = procs[idx]
                ppt = cell_des.add_paragraph()
                ppt.paragraph_format.space_after = Pt(4)
                rpt = ppt.add_run(proc.titulo.upper())
                rpt.bold = True
                rpt.font.size = Pt(9.5)
                rpt.font.color.rgb = RGBColor(192, 57, 43)
                for par in proc.contenido:
                    ppr = cell_des.add_paragraph()
                    ppr.paragraph_format.space_after = Pt(4)
                    ppr.paragraph_format.line_spacing = 1.15
                    ppr.add_run(par).font.size = Pt(9.5)
            else:
                cell_des.add_paragraph().add_run("Gestion y Acompanamiento del Desarrollo de Competencias...").font.size = Pt(9.5)

            if idx > 0:
                _write_vertical_cell(mt.cell(rn, 1), "MOTIVACION")
                _write_vertical_cell(mt.cell(rn, 3), "EVALUACION")

        if n_proc > 1:
            mt.cell(2, 0).merge(mt.cell(2 + n_proc - 1, 0))
            mt.cell(2, 1).merge(mt.cell(2 + n_proc - 1, 1))
            mt.cell(2, 3).merge(mt.cell(2 + n_proc - 1, 3))

        # Fila CIERRE
        rc = n_rows - 1
        _write_momento_cell(mt.cell(rc, 0), "CIERRE",
            "Evaluacion (Reflexion sobre lo aprendido)\nAcciones de reforzamiento o indagacion",
            session.momentos.cierre.tiempo_total or "")
        _write_vertical_cell(mt.cell(rc, 1), "MOTIVACION")
        _write_vertical_cell(mt.cell(rc, 3), "EVALUACION")

        cell_cie = mt.cell(rc, 2)
        set_cell_margins(cell_cie, top=120, bottom=120, left=140, right=140)
        cell_cie.text = ""

        cierre = session.momentos.cierre
        cierre_has_content = bool(cierre.metacognicion or cierre.evaluacion or cierre.extension)

        if cierre.metacognicion:
            plbl = cell_cie.add_paragraph()
            rlbl = plbl.add_run("Metacognicion:")
            rlbl.bold = True
            rlbl.font.size = Pt(9.5)
            for m in cierre.metacognicion:
                pi = cell_cie.add_paragraph(style='List Bullet')
                pi.paragraph_format.space_after = Pt(2)
                pi.add_run(m).font.size = Pt(9.5)
        if cierre.evaluacion:
            plbl = cell_cie.add_paragraph()
            plbl.paragraph_format.space_before = Pt(4)
            rlbl = plbl.add_run("Evaluacion formativa:")
            rlbl.bold = True
            rlbl.font.size = Pt(9.5)
            for ev in cierre.evaluacion:
                pi = cell_cie.add_paragraph(style='List Bullet')
                pi.paragraph_format.space_after = Pt(2)
                pi.add_run(ev).font.size = Pt(9.5)
        if cierre.extension:
            plbl = cell_cie.add_paragraph()
            plbl.paragraph_format.space_before = Pt(4)
            rlbl = plbl.add_run("Extension para casa:")
            rlbl.bold = True
            rlbl.font.size = Pt(9.5)
            for ext in cierre.extension:
                pi = cell_cie.add_paragraph(style='List Bullet')
                pi.paragraph_format.space_after = Pt(2)
                pi.add_run(ext).font.size = Pt(9.5)

        if not cierre_has_content:
            default_sections = [
                ("Metacognicion:", [
                    "Que aprendimos hoy? Como lo aprendimos? Para que nos sirve?",
                    "Que fue lo mas dificil? Como lo superamos?"
                ]),
                ("Evaluacion formativa:", [
                    "Revision de los criterios de evaluacion con los estudiantes.",
                    "Retroalimentacion sobre el desempeno de la sesion."
                ]),
                ("Extension para casa:", [
                    "Actividad de refuerzo o aplicacion a nuevas situaciones.",
                    "Resolucion de ejercicios complementarios."
                ]),
            ]
            for lbl, items in default_sections:
                plbl = cell_cie.add_paragraph()
                plbl.paragraph_format.space_before = Pt(4)
                rlbl = plbl.add_run(lbl)
                rlbl.bold = True
                rlbl.font.size = Pt(9.5)
                for item in items:
                    pi = cell_cie.add_paragraph(style='List Bullet')
                    pi.paragraph_format.space_after = Pt(2)
                    pi.add_run(item).font.size = Pt(9.5)

        anchos_mom = [Inches(1.2), Inches(0.35), Inches(4.87), Inches(0.35)]
        for row in mt.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = anchos_mom[ci]

    # ===============================================================
    # FIRMAS DE LA SESION
    # ===============================================================
    doc.add_paragraph().paragraph_format.space_before = Pt(40)
    ft = doc.add_table(rows=1, cols=2)
    ft.autofit = True
    for cell in ft.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcB.append(b)
        tcPr.append(tcB)
    pfd = ft.cell(0, 0).paragraphs[0]
    pfd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pfd.add_run("_______________________________\n").bold = True
    rnd = pfd.add_run((session.metadata.docente or "Docente de la Sesion") + "\n")
    rnd.bold = True
    rnd.font.size = Pt(9.5)
    rcd = pfd.add_run("Docente de la Sesion")
    rcd.font.size = Pt(8.5)

    # Firma Director
    pfs = ft.cell(0, 1).paragraphs[0]
    pfs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pfs.add_run("_______________________________\n").bold = True
    rns = pfs.add_run((session.metadata.director or "Director de la I.E.") + "\n")
    rns.bold = True
    rns.font.size = Pt(9.5)
    rcs = pfs.add_run("Director de la I.E.")
    rcs.font.size = Pt(8.5)

    # ===============================================================
    # FICHA DE TRABAJO (SI SE SOLICITA)
    # ===============================================================
    if session.ficha_trabajo and session.ficha_trabajo.titulo:
        doc.add_page_break()
        
        f_sec = doc.add_section()
        f_sec.top_margin = Inches(0.8)
        f_sec.bottom_margin = Inches(0.8)
        f_sec.left_margin = Inches(0.8)
        f_sec.right_margin = Inches(0.8)

        ft_tbl = doc.add_table(rows=1, cols=1)
        ft_tbl.autofit = False
        ft_tbl.rows[0].cells[0].width = Inches(6.67)
        add_table_borders_black(ft_tbl)
        set_cell_text_white_bold(ft_tbl.cell(0, 0), "FICHA DE TRABAJO DE APRENDIZAJE INDEPENDIENTE", font_size_pt=11.5)
        set_cell_background(ft_tbl.cell(0, 0), '2980B9')
        set_cell_margins(ft_tbl.cell(0, 0), top=120, bottom=120, left=180, right=180)

        doc.add_paragraph().paragraph_format.space_before = Pt(10)

        stud_tbl = doc.add_table(rows=1, cols=2)
        stud_tbl.autofit = False
        add_table_borders(stud_tbl, color='CBD5E1', sz='4')
        
        set_cell_margins(stud_tbl.cell(0, 0), top=80, bottom=80, left=120, right=120)
        p_st1 = stud_tbl.cell(0, 0).paragraphs[0]
        r_st1 = p_st1.add_run("Estudiante: __________________________________________________")
        r_st1.bold = True
        r_st1.font.size = Pt(9.5)
        
        set_cell_margins(stud_tbl.cell(0, 1), top=80, bottom=80, left=120, right=120)
        p_st2 = stud_tbl.cell(0, 1).paragraphs[0]
        p_st2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_st2 = p_st2.add_run("Grado y Sección: ________________")
        r_st2.bold = True
        r_st2.font.size = Pt(9.5)

        for row in stud_tbl.rows:
            row.cells[0].width = Inches(4.5)
            row.cells[1].width = Inches(2.17)

        doc.add_paragraph().paragraph_format.space_before = Pt(8)

        p_ft_t = doc.add_paragraph()
        p_ft_t.paragraph_format.space_after = Pt(4)
        rf_t = p_ft_t.add_run("Actividad: " + session.ficha_trabajo.titulo.upper())
        rf_t.bold = True
        rf_t.font.size = Pt(12)
        rf_t.font.color.rgb = RGBColor(41, 128, 185)

        if session.ficha_trabajo.indicaciones:
            p_ft_ind = doc.add_paragraph()
            p_ft_ind.paragraph_format.space_after = Pt(12)
            p_ft_ind.paragraph_format.line_spacing = 1.15
            p_ft_ind.add_run("Indicaciones: ").bold = True
            p_ft_ind.runs[0].font.size = Pt(9.5)
            p_ft_ind.add_run(session.ficha_trabajo.indicaciones).font.size = Pt(9.5)
            p_ft_ind.runs[1].font.italic = True

        act_html = session.ficha_trabajo.actividades or ""
        soup_act = BeautifulSoup(act_html, 'html.parser')
        
        def add_act_element(element, is_bold=False, is_italic=False, list_style=None):
            if isinstance(element, Tag):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(4)
                    r = p.add_run(element.get_text().strip())
                    r.bold = True
                    r.font.size = Pt(11)
                elif element.name == 'p':
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    add_runs_to_paragraph(p, element, is_bold, is_italic)
                elif element.name == 'ul':
                    for li in element.find_all('li', recursive=False):
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.space_after = Pt(3)
                        p.paragraph_format.line_spacing = 1.15
                        add_runs_to_paragraph(p, li, is_bold, is_italic)
                elif element.name == 'ol':
                    for li in element.find_all('li', recursive=False):
                        p = doc.add_paragraph(style='List Number')
                        p.paragraph_format.space_after = Pt(3)
                        p.paragraph_format.line_spacing = 1.15
                        add_runs_to_paragraph(p, li, is_bold, is_italic)
                elif element.name == 'table':
                    rows = element.find_all('tr')
                    if rows:
                        max_c = max(len(r.find_all(['td', 'th'])) for r in rows)
                        if max_c > 0:
                            tbl = doc.add_table(rows=len(rows), cols=max_c)
                            tbl.autofit = True
                            add_table_borders(tbl)
                            for ri, row in enumerate(rows):
                                cells = row.find_all(['td', 'th'])
                                for ci, cell_h in enumerate(cells):
                                    if ci < max_c:
                                        c = tbl.cell(ri, ci)
                                        set_cell_margins(c, top=80, bottom=80, left=100, right=100)
                                        p_cell = c.paragraphs[0]
                                        p_cell.paragraph_format.space_after = Pt(0)
                                        p_cell.paragraph_format.line_spacing = 1.1
                                        if cell_h.name == 'th':
                                            set_cell_background(c, 'F2F2F2')
                                            add_runs_to_paragraph(p_cell, cell_h, is_bold=True)
                                        else:
                                            add_runs_to_paragraph(p_cell, cell_h)
                else:
                    for child in element.children:
                        add_act_element(child, is_bold, is_italic)

        for child in soup_act.children:
            add_act_element(child)

    # ===============================================================
    # TABLA LISTA DE COTEJO (SI SE INCLUYEN ALUMNOS)
    # ===============================================================
    alumnos = session.alumnos
    criterios = session.proposito.criterios
    
    if alumnos and len(alumnos) > 0 and len(criterios) > 0:
        doc.add_page_break()
        
        lc_sec = doc.add_section()
        lc_sec.page_width = Inches(11.69)
        lc_sec.page_height = Inches(8.27)
        lc_sec.top_margin = Inches(0.6)
        lc_sec.bottom_margin = Inches(0.6)
        lc_sec.left_margin = Inches(0.6)
        lc_sec.right_margin = Inches(0.6)
        
        lc_tbl = doc.add_table(rows=1, cols=1)
        lc_tbl.autofit = False
        lc_tbl.rows[0].cells[0].width = Inches(10.49)
        add_table_borders_black(lc_tbl)
        set_cell_text_white_bold(lc_tbl.cell(0, 0), "LISTA DE COTEJO DE EVALUACION FORMATIVA", font_size_pt=12)
        set_cell_background(lc_tbl.cell(0, 0), '2C3E50')
        set_cell_margins(lc_tbl.cell(0, 0), top=120, bottom=120, left=180, right=180)

        doc.add_paragraph().paragraph_format.space_before = Pt(8)

        lch_tbl = doc.add_table(rows=1, cols=4)
        lch_tbl.autofit = False
        add_table_borders_black(lch_tbl)
        
        _label(lch_tbl.cell(0, 0), "IE / Area")
        _val(lch_tbl.cell(0, 1), f"{session.metadata.institucion or 'IE'} / {session.metadata.area or 'Matematica'}")
        _label(lch_tbl.cell(0, 2), "Grado / Seccion")
        _val(lch_tbl.cell(0, 3), f"{session.metadata.grado or ''} \"{session.metadata.seccion or ''}\"")

        for row in lch_tbl.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(4.5)
            row.cells[2].width = Inches(1.5)
            row.cells[3].width = Inches(2.99)

        doc.add_paragraph().paragraph_format.space_before = Pt(8)

        num_cols = 2 + len(criterios) * 2
        lct = doc.add_table(rows=2 + len(alumnos), cols=num_cols)
        lct.autofit = False
        add_table_borders_black(lct)

        lct.cell(0, 0).merge(lct.cell(1, 0))
        lct.cell(0, 1).merge(lct.cell(1, 1))

        for ci, crit in enumerate(criterios):
            sc = 2 + ci * 2
            lct.cell(0, sc).merge(lct.cell(0, sc + 1))
            lct.cell(0, sc).paragraphs[0].text = f"Criterio {ci + 1}: {crit}"
            lct.cell(1, sc).paragraphs[0].text = "SI"
            lct.cell(1, sc + 1).paragraphs[0].text = "NO"

        CRIT_COLORS   = ['D9E1F2', 'FADBD8', 'D5F5E3', 'FCF3CF', 'FDE8D8', 'E8DAEF']
        SUBCRIT_COLORS = ['BDD7EE', 'FADBD8', 'A9DFBF', 'F9E79F', 'FAD7A0', 'D7BDE2']

        def _lcfmt(cell, width_in, font_size_pt, bold=False, ctr=False, bg=None):
            cell.width = Inches(width_in)
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            if bg:
                set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if ctr:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = bold
                run.font.size = Pt(font_size_pt)

        _lcfmt(lct.cell(0, 0), 0.35, 8.5, bold=True, ctr=True, bg='FFF2CC')
        _lcfmt(lct.cell(0, 1), 2.2,  8.5, bold=True, ctr=False, bg='FFF2CC')
        for ci, crit in enumerate(criterios):
            sc = 2 + ci * 2
            _lcfmt(lct.cell(0, sc),     0.7,  7.5, bold=True, ctr=True, bg=CRIT_COLORS[ci % len(CRIT_COLORS)])
            _lcfmt(lct.cell(1, sc),     0.35, 8,   bold=True, ctr=True, bg=SUBCRIT_COLORS[ci % len(SUBCRIT_COLORS)])
            _lcfmt(lct.cell(1, sc + 1), 0.35, 8,   bold=True, ctr=True, bg=SUBCRIT_COLORS[ci % len(SUBCRIT_COLORS)])
        for ri, stud in enumerate(alumnos):
            rn = 2 + ri
            lct.cell(rn, 0).text = str(ri + 1)
            lct.cell(rn, 1).text = "" if stud.startswith("Estudiante ") else stud
            _lcfmt(lct.cell(rn, 0), 0.35, 8.5, bold=True,  ctr=True)
            _lcfmt(lct.cell(rn, 1), 2.2,  8.5, bold=False, ctr=False)
            for ci in range(len(criterios)):
                sc = 2 + ci * 2
                _lcfmt(lct.cell(rn, sc),     0.35, 8, ctr=True)
                _lcfmt(lct.cell(rn, sc + 1), 0.35, 8, ctr=True)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream
