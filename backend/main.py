import os
import io
import re
import sys
import secrets
import asyncio
import webbrowser
import threading
import subprocess
import warnings
import urllib.request
import zipfile
import tkinter as tk
from tkinter import scrolledtext
from pathlib import Path

# Determinar base absoluta del ejecutable para portabilidad
if getattr(sys, 'frozen', False):
    EXE_DIR = Path(sys.executable).parent
    BASE_DIR = Path(sys._MEIPASS)
else:
    EXE_DIR = Path(__file__).resolve().parent
    BASE_DIR = EXE_DIR.parent

LOCAL_BIN_DIR = EXE_DIR / "bin"
CHROMIUM_DIR = LOCAL_BIN_DIR / "chrome-win"
CHROMIUM_EXE = CHROMIUM_DIR / "chrome.exe"

# Desactivar advertencias molestas en la consola (como DeprecationWarnings de FastAPI/Lifespan)
warnings.filterwarnings("ignore")

# Configurar consola en Windows para UTF-8 de forma forzada para evitar fallos con emojis/bloques
if sys.platform.startswith('win'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except Exception:
        pass

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel
from playwright.async_api import async_playwright
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Librerías para estilizar consola
try:
    from rich import print as rprint
    from rich.console import Console
    from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn
    from rich.panel import Panel
    console = Console()
except ImportError:
    # Fallback si no está instalado rich
    console = None
    def rprint(*args, **kwargs):
        print(*args, **kwargs)

# Inicialización de FastAPI
app = FastAPI(
    title="Motor de Exportación Pedagógica",
    description="Backend local para generación premium de PDFs y Word (.docx)"
)

# Configuración de CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.middleware("http")
async def add_private_network_header(request, call_next):
    if request.method == "OPTIONS":
        response = Response()
        response.headers["Access-Control-Allow-Origin"] = request.headers.get("origin", "*")
        response.headers["Access-Control-Allow-Methods"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "*"
        response.headers["Access-Control-Allow-Private-Network"] = "true"
        return response

    response = await call_next(request)
    response.headers["Access-Control-Allow-Private-Network"] = "true"
    return response

# Variables globales para el enlace de sesión
CONNECTION_TOKEN = secrets.token_hex(16)
CLIENT_CONNECTED = False

# Esquemas de Datos con Token obligatorio para mayor seguridad
class ExportPDFRequest(BaseModel):
    html_content: str
    titulo: str = "Sesion_de_Aprendizaje"
    token: str

class ExportDocxRequest(BaseModel):
    html_content: str
    titulo: str = "Sesion_de_Aprendizaje"
    token: str

# ── MODELO DE DATOS JSON ESTRUCTURADO PARA SESIONES PREMIUM ──
from typing import List, Optional

class MetadataData(BaseModel):
    institucion: Optional[str] = ""
    dre: Optional[str] = ""
    ugel: Optional[str] = ""
    docente: Optional[str] = ""
    director: Optional[str] = ""
    fecha: Optional[str] = ""
    nivel: Optional[str] = ""
    numero_sesion: Optional[str] = ""
    grado: Optional[str] = ""
    seccion: Optional[str] = ""
    area: Optional[str] = ""
    duracion: Optional[str] = ""
    unidad: Optional[str] = ""
    titulo: Optional[str] = ""
    logo_left_url: Optional[str] = ""
    logo_regional_url: Optional[str] = ""

class PropositoData(BaseModel):
    proposito_texto: Optional[str] = ""
    conocimientos: Optional[str] = ""
    competencia: Optional[str] = ""
    estandar: Optional[str] = ""
    capacidades: List[str] = []
    criterios: List[str] = []
    producto_evidencia: Optional[str] = ""
    instrumento: Optional[str] = ""

class CompetenciaTransversal(BaseModel):
    titulo: str
    desempenos: List[str] = []

class EnfoqueTransversal(BaseModel):
    nombre: str
    valor: str
    actitudes: str

class RecursosData(BaseModel):
    enlaces: Optional[str] = ""
    materiales: Optional[str] = ""
    refuerzo: Optional[str] = ""

class MomentoInicio(BaseModel):
    tiempo_total: Optional[str] = ""
    actividades: List[str] = []

class ProcesoDesarrollo(BaseModel):
    clave: str
    titulo: str
    contenido: List[str] = []

class MomentoDesarrollo(BaseModel):
    tiempo_total: Optional[str] = ""
    procesos: List[ProcesoDesarrollo] = []

class MomentoCierre(BaseModel):
    tiempo_total: Optional[str] = ""
    metacognicion: List[str] = []
    evaluacion: List[str] = []
    extension: List[str] = []

class MomentosData(BaseModel):
    inicio: MomentoInicio
    desarrollo: MomentoDesarrollo
    cierre: MomentoCierre

class FichaTrabajoData(BaseModel):
    titulo: Optional[str] = ""
    indicaciones: Optional[str] = ""
    actividades: Optional[str] = ""

class SesionAprendizajeRequest(BaseModel):
    metadata: MetadataData
    proposito: PropositoData
    competencias_transversales: List[CompetenciaTransversal] = []
    enfoques_transversales: List[EnfoqueTransversal] = []
    recursos: RecursosData
    momentos: MomentosData
    ficha_trabajo: Optional[FichaTrabajoData] = None
    alumnos: Optional[List[str]] = []
    token: str


@app.get("/")
def check_status():
    """Endpoint de control para verificar si el servidor local está activo."""
    return {
        "status": "Online",
        "engine": "FastAPI + Python Export Engine",
        "connected": CLIENT_CONNECTED
    }


@app.get("/verificar-token")
def verificar_token(token: str):
    """Verifica si el token proveído coincide con el de la sesión actual."""
    global CLIENT_CONNECTED
    if token == CONNECTION_TOKEN:
        if not CLIENT_CONNECTED:
            CLIENT_CONNECTED = True
            print("\n⚡ [CONEXIÓN ESTABLECIDA] El navegador se ha enlazado con éxito.\n")
        return {"status": "Connected", "message": "Enlace establecido correctamente."}
    else:
        raise HTTPException(status_code=401, detail="Token de conexión inválido.")


@app.post("/exportar-pdf")
async def exportar_pdf(payload: ExportPDFRequest):
    """
    Exporta el HTML y CSS recibido a un archivo PDF físico A4
    utilizando Playwright (Chromium headless) con paginado dinámico y reglas anti-corte de tablas.
    """
    if payload.token != CONNECTION_TOKEN:
        raise HTTPException(status_code=401, detail="No autorizado: Token de conexión inválido.")

    try:
        # Sanitizar el nombre del archivo
        filename = re.sub(r'[^a-zA-Z0-9-_\s]', '', payload.titulo).replace(' ', '_')
        nombre_archivo = f"{filename}.pdf"

        # HTML base mínimo — el frontend ya trae el HTML pre-paginado en divs .hoja-a4
        documento_completo = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{payload.titulo}</title>
            <style>
                /* Glue CSS mínimo: el frontend es responsable del diseño y la paginación */
                * {{ box-sizing: border-box; margin: 0; padding: 0; }}
                body {{
                    background: #fff;
                    -webkit-print-color-adjust: exact;
                    print-color-adjust: exact;
                }}
                /* Garantizar salto de página después de cada hoja */
                .hoja-a4 {{
                    page-break-after: always !important;
                    break-after: page !important;
                }}
                /* Ocultar elementos interactivos residuales */
                .no-print, .add-logo-placeholder, .btn-remove-logo {{
                    display: none !important;
                }}
            </style>
        </head>
        <body>
            {payload.html_content}
        </body>
        </html>
        """

        async with async_playwright() as p:
            ruta_motor = buscar_navegador_compatible()
            launch_args = {"headless": True}
            if ruta_motor:
                launch_args["executable_path"] = ruta_motor
            browser = await p.chromium.launch(**launch_args)
            context = await browser.new_context()
            page = await context.new_page()
            
            await page.set_content(documento_completo, wait_until="networkidle")
            
            # Usar prefer_css_page_size=True → Chromium respeta la medida exacta del div .hoja-a4
            # Los márgenes van en cero porque el padding ya está definido en el div
            pdf_bytes = await page.pdf(
                print_background=True,
                prefer_css_page_size=True,
                margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
                display_header_footer=False
            )
            await browser.close()

        if console:
            console.print(f"[green]✓ [PDF EXPORTADO] Generado con éxito: {nombre_archivo}[/green]")

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={nombre_archivo}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )

    except Exception as e:
        print("[ERROR PDF]", str(e))
        raise HTTPException(status_code=500, detail=f"Fallo al compilar PDF: {str(e)}")


def escape_html(text: str) -> str:
    """Escapa caracteres HTML básicos."""
    if not text:
        return ""
    return (text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
                .replace("'", "&#x27;"))


def build_pdf_html_from_json(session: SesionAprendizajeRequest) -> str:
    # 1. Cabecera con logos si existen
    logo_left_html = ""
    if session.metadata.logo_left_url:
        logo_left_html = f'<img src="{session.metadata.logo_left_url}" class="header-logo-img" />'

    logo_right_html = ""
    if session.metadata.logo_regional_url:
        logo_right_html = f'<img src="{session.metadata.logo_regional_url}" class="header-logo-img" />'

    # 2. Listas de propósitos
    capacidades_html = "".join([f"<li>{escape_html(c)}</li>" for c in session.proposito.capacidades])
    criterios_html = "".join([f"<li>{escape_html(c)}</li>" for c in session.proposito.criterios])

    # 3. Competencias transversales
    ct_rows_html = ""
    if session.competencias_transversales:
        for ct in session.competencias_transversales:
            desempenos_li = "".join([f"<li>{escape_html(d)}</li>" for d in ct.desempenos])
            ct_rows_html += f"""
            <tr>
                <td style="font-weight: 600;">{escape_html(ct.titulo)}</td>
                <td><ul class="session-list">{desempenos_li}</ul></td>
            </tr>
            """

    # 4. Enfoques transversales
    enfoques_rows_html = ""
    if session.enfoques_transversales:
        for enf in session.enfoques_transversales:
            enfoques_rows_html += f"""
            <tr>
                <td style="font-weight: 600;">{escape_html(enf.nombre)}</td>
                <td>{escape_html(enf.valor)}</td>
                <td>{escape_html(enf.actitudes)}</td>
            </tr>
            """

    # 5. Momentos Didácticos (Fusión inteligente con rowspan en HTML)
    procesos_des = session.momentos.desarrollo.processes if hasattr(session.momentos.desarrollo, 'processes') else session.momentos.desarrollo.procesos
    cant_procesos = len(procesos_des) if procesos_des else 1

    # Inicio
    inicio_actividades_html = "".join([f"<p class='proceso-parrafo'>{escape_html(act)}</p>" for act in session.momentos.inicio.actividades])
    
    # Desarrollo (Primer proceso y siguientes)
    desarrollo_primero_html = ""
    desarrollo_siguientes_html = ""
    
    if procesos_des:
        p_primero = procesos_des[0]
        p_primero_cont = "".join([f"<p class='proceso-parrafo'>{escape_html(par)}</p>" for par in p_primero.contenido])
        desarrollo_primero_html = f"""
        <div class="proceso-titulo">{escape_html(p_primero.titulo)}</div>
        {p_primero_cont}
        """
        
        for idx in range(1, cant_procesos):
            p_sig = procesos_des[idx]
            p_sig_cont = "".join([f"<p class='proceso-parrafo'>{escape_html(par)}</p>" for par in p_sig.contenido])
            desarrollo_siguientes_html += f"""
            <tr>
                <td>
                    <div class="proceso-titulo">{escape_html(p_sig.titulo)}</div>
                    {p_sig_cont}
                </td>
            </tr>
            """
    else:
        desarrollo_primero_html = "<p class='proceso-parrafo'>Gestión y Acompañamiento del Desarrollo de Competencias...</p>"

    # Cierre
    cierre_estrategias_html = ""
    if session.momentos.cierre.metacognicion:
        cierre_estrategias_html += "<p class='proceso-parrafo'><strong>Metacognición:</strong></p><ul class='session-list'>"
        cierre_estrategias_html += "".join([f"<li>{escape_html(m)}</li>" for m in session.momentos.cierre.metacognicion])
        cierre_estrategias_html += "</ul>"
    if session.momentos.cierre.evaluacion:
        cierre_estrategias_html += "<p class='proceso-parrafo' style='margin-top:8px;'><strong>Evaluación formativa:</strong></p><ul class='session-list'>"
        cierre_estrategias_html += "".join([f"<li>{escape_html(e)}</li>" for e in session.momentos.cierre.evaluacion])
        cierre_estrategias_html += "</ul>"
    if session.momentos.cierre.extension:
        cierre_estrategias_html += "<p class='proceso-parrafo' style='margin-top:8px;'><strong>Extensión para casa:</strong></p><ul class='session-list'>"
        cierre_estrategias_html += "".join([f"<li>{escape_html(ext)}</li>" for ext in session.momentos.cierre.extension])
        cierre_estrategias_html += "</ul>"

    # Ficha de Trabajo
    ficha_html = ""
    if session.ficha_trabajo:
        ficha_actividades_txt = session.ficha_trabajo.actividades or ""
        ficha_actividades_txt = re.sub(r'<[^>]*>', '', ficha_actividades_txt) # Limpiar tags HTML residuales
        ficha_actividades_p = "".join([f"<p class='proceso-parrafo'>{escape_html(p)}</p>" for p in ficha_actividades_txt.split("\n") if p.strip()])
        
        ficha_html = f"""
        <div class="hoja-a4" style="page-break-before: always; break-before: page;">
            <div class="ficha-title">FICHA DE TRABAJO INDEPENDIENTE PARA EL ESTUDIANTE</div>
            
            <table class="ficha-header-table">
                <tr>
                    <td>Nombre: __________________________________________________</td>
                    <td style="text-align: right;">Grado y Sección: ________________</td>
                </tr>
            </table>
            
            <div class="ficha-act-title">🎨 Actividad: {escape_html(session.ficha_trabajo.titulo or 'Mi Ficha Práctica')}</div>
            <div class="ficha-indicaciones">
                <strong>Indicaciones: </strong><span>{escape_html(session.ficha_trabajo.indicaciones or 'Realiza la actividad según las indicaciones.')}</span>
            </div>
            
            <div class="ficha-contenido">
                {ficha_actividades_p}
            </div>
        </div>
        """

    # 5.5. Construir la Lista de Cotejo dinámica basada en la lista de alumnos
    alumnos_list = session.alumnos if session.alumnos else []
    if not alumnos_list:
        alumnos_list = [f"Estudiante {i+1}" for i in range(30)]

    criterios = session.proposito.criterios if session.proposito.criterios else []
    if not criterios:
        criterios = [
            "Expresa con diversas representaciones la comprensión sobre el tema.",
            "Ordena y organiza conceptos clave para resolver problemas.",
            "Emplea estrategias y procedimientos diversos para realizar las tareas.",
            "Halla y valida soluciones utilizando criterios y conocimientos del área."
        ]

    criterios_headers_html = "".join([
        f"<th colspan='2' style='font-size: 7.5pt; font-weight: bold; background: #e2e8f0; border: 1px solid #000; padding: 4px; text-align: center; vertical-align: top; max-width: 150px;'>{escape_html(c)}</th>"
        for c in criterios
    ])

    criterios_subheaders_html = "".join([
        "<th style='width: 30px; text-align: center; background: #f1f5f9; border: 1px solid #000; font-size: 8pt; font-weight: bold;'>SI</th><th style='width: 30px; text-align: center; background: #f1f5f9; border: 1px solid #000; font-size: 8pt; font-weight: bold;'>NO</th>"
        for _ in criterios
    ])

    rows_html = ""
    for idx, stud in enumerate(alumnos_list):
        display_name = "" if stud.startswith("Estudiante ") else stud
        criterios_cells_html = "".join([
            "<td style='border: 1px solid #000;'></td><td style='border: 1px solid #000;'></td>"
            for _ in criterios
        ])
        rows_html += f"""
        <tr>
            <td style="text-align: center; font-weight: 700; height: 26px; border: 1px solid #000; font-size: 8.5pt;">{idx + 1}</td>
            <td style="text-align: left; padding-left: 6px; font-weight: 600; border: 1px solid #000; font-size: 8.5pt;">{escape_html(display_name)}</td>
            {criterios_cells_html}
        </tr>
        """

    lista_cotejo_html = f"""
    <div class="hoja-a4" style="page-break-before: always; break-before: page; padding: 18mm 12mm;">
        <div class="section-title" style="text-align: center; font-size: 11pt; font-weight: 800; text-transform: uppercase;">Instrumento de Evaluación</div>
        <div class="section-title" style="text-align: center; font-size: 9.5pt; font-weight: 700; background: #f1f5f9; color: #000; margin-top: 4px;">
            LISTA DE COTEJO {escape_html(session.metadata.grado or '2°')} {escape_html(session.metadata.seccion or 'A')}
        </div>
        
        <table class="content-table momentos-table" style="width: 100%; border-collapse: collapse; margin-top: 10px; border: 1px solid #000;">
            <thead>
                <tr>
                    <th rowspan="2" style="width: 30px; text-align: center; background: #e2e8f0; border: 1px solid #000; font-size: 8.5pt;">N°</th>
                    <th rowspan="2" style="text-align: left; padding-left: 6px; background: #e2e8f0; border: 1px solid #000; font-size: 8.5pt;">ESTUDIANTES</th>
                    {criterios_headers_html}
                </tr>
                <tr>
                    {criterios_subheaders_html}
                </tr>
            </thead>
            <tbody>
                {rows_html}
            </tbody>
        </table>
    </div>
    """

    # 6. HTML final unificado
    html_content = f"""<!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <!-- Cargar MathJax para renderizar ecuaciones matemáticas en el PDF -->
        <script>
            window.MathJax = {{
                tex: {{
                    inlineMath: [['$', '$'], ['\\(', '\\)']],
                    displayMath: [['$$', '$$'], ['\\[', '\\]']],
                    processEscapes: true
                }},
                options: {{
                    ignoreHtmlClass: 'tex2jax_ignore',
                    processHtmlClass: 'tex2jax_process'
                }},
                svg: {{
                    fontCache: 'global'
                }}
            }};
        </script>
        <script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-svg.js"></script>
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');
            
            * {{ box-sizing: border-box; margin: 0; padding: 0; }}
            body {{
                font-family: 'Inter', Arial, sans-serif;
                background: #ffffff;
                color: #1e293b;
                font-size: 10pt;
                line-height: 1.4;
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }}

            @page {
                size: A4 portrait;
                margin: 0 !important;
            }

            .hoja-a4 {
                width: 210mm;
                height: 297mm;
                min-height: 297mm;
                max-height: 297mm;
                padding: 18mm 12mm;
                box-sizing: border-box;
                background: #ffffff;
                position: relative;
                page-break-after: always;
                break-after: page;
                overflow: hidden;
            }

            /* Cabecera */
            .header-table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 8px;
            }}
            .header-table td {{
                border: none !important;
                padding: 0;
                vertical-align: middle;
            }}
            .header-logos {{
                width: 70px;
            }}
            .header-logo-img {{
                max-width: 65px;
                max-height: 65px;
                object-fit: contain;
            }}
            .header-text {{
                text-align: center;
                font-size: 7.5pt;
                line-height: 1.25;
                color: #334155;
            }}
            .header-text .minedu {{
                font-weight: 700;
                font-size: 8pt;
                color: #0f172a;
            }}
            .header-text .dre, .header-text .ugel {{
                font-weight: 600;
                font-size: 8pt;
            }}
            .header-text .agp {{
                font-style: italic;
                color: #64748b;
            }}

            .divider {{
                border-bottom: 2px solid #0f172a;
                margin-top: 4px;
                margin-bottom: 12px;
            }}

            .title-box {{
                text-align: center;
                margin-bottom: 12px;
            }}
            .title-box h1 {{
                font-size: 13pt;
                font-weight: 700;
                margin: 0;
                color: #0f172a;
                text-transform: uppercase;
            }}
            .title-box h2 {{
                font-size: 10.5pt;
                font-weight: 600;
                font-style: italic;
                margin: 2px 0 0 0;
                color: #334155;
            }}

            /* Tablas */
            table.content-table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 10px;
                page-break-inside: auto;
            }}
            table.content-table tr {{
                page-break-inside: avoid !important;
                break-inside: avoid !important;
            }}
            table.content-table th, table.content-table td {{
                border: 1px solid #cbd5e1;
                padding: 5px 7px;
                font-size: 9pt;
                vertical-align: top;
            }}
            table.content-table th {{
                background-color: #f1f5f9;
                font-weight: 600;
                text-align: left;
                color: #0f172a;
            }}
            table.content-table td.label-cell {{
                background-color: #f8fafc;
                font-weight: 600;
                color: #334155;
                width: 140px;
            }}

            .section-title {{
                font-size: 10.5pt;
                font-weight: 700;
                color: #0f172a;
                margin: 12px 0 4px 0;
                text-transform: uppercase;
                border-left: 3px solid #3b82f6;
                padding-left: 6px;
            }}

            .section-content {{
                font-size: 9.5pt;
                color: #334155;
                padding-left: 4px;
                margin-bottom: 10px;
            }}

            /* Listas */
            ul.session-list {{
                margin: 0;
                padding-left: 14px;
            }}
            ul.session-list li {{
                margin-bottom: 2px;
            }}

            /* Momentos didácticos */
            .momentos-table th {{
                background-color: #e2e8f0 !important;
                font-size: 8.5pt !important;
                text-align: center !important;
            }}
            .momento-label-cell {{
                background-color: #f8fafc;
                font-weight: 700;
                font-size: 9pt;
                color: #0f172a;
                width: 120px;
            }}
            .momento-time {{
                font-size: 7.5pt;
                color: #64748b;
                margin-top: 4px;
                font-weight: 600;
            }}
            .momento-eval-cell {{
                background-color: #f8fafc;
                font-weight: 600;
                font-size: 8pt;
                color: #475569;
                width: 95px;
            }}
            .proceso-titulo {{
                font-weight: 700;
                color: #b91c1c; /* C0392B */
                font-size: 8.5pt;
                text-transform: uppercase;
                margin-bottom: 4px;
            }}
            .proceso-parrafo {{
                margin: 0 0 4px 0;
                font-size: 9pt;
            }}

            /* Firmas */
            .firmas-container {{
                display: flex;
                justify-content: space-between;
                margin-top: 35px;
                padding: 0 30px;
                page-break-inside: avoid !important;
                break-inside: avoid !important;
            }}
            .firma-box {{
                text-align: center;
                width: 180px;
            }}
            .firma-linea {{
                border-top: 1px solid #475569;
                margin-bottom: 4px;
            }}
            .firma-nombre {{
                font-weight: 600;
                font-size: 8.5pt;
                color: #0f172a;
            }}
            .firma-cargo {{
                font-size: 7.5pt;
                color: #64748b;
            }}

            /* Ficha de trabajo */
            .ficha-title {{
                text-align: center;
                font-size: 11pt;
                font-weight: 700;
                color: #0f172a;
                margin-top: 5px;
                margin-bottom: 15px;
                text-transform: uppercase;
            }}
            .ficha-header-table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 12px;
            }}
            .ficha-header-table td {{
                border-bottom: 2px solid #3498db;
                padding: 5px 0;
                font-weight: 600;
                font-size: 9pt;
                color: #2c3e50;
            }}
            .ficha-act-title {{
                font-size: 10.5pt;
                font-weight: 700;
                color: #2980b9;
                margin-top: 12px;
                margin-bottom: 4px;
            }}
            .ficha-indicaciones {{
                font-size: 9pt;
                margin-bottom: 10px;
            }}
            .ficha-indicaciones strong {{
                color: #0f172a;
            }}
            .ficha-indicaciones span {{
                font-style: italic;
                color: #555;
            }}
            .ficha-contenido {{
                font-size: 9pt;
                color: #334155;
                white-space: pre-wrap;
            }}
        </style>
    </head>
    <body>
        <div class="hoja-a4">
            <!-- ════════ CABECERA INSTITUCIONAL ════════ -->
            <table class="header-table">
                <tr>
                    <td class="header-logos" style="text-align: left;">
                        {logo_left_html}
                    </td>
                    <td class="header-text">
                        <span class="minedu">MINISTERIO DE EDUCACIÓN</span><br>
                        <span class="dre">{escape_html(session.metadata.dre) or 'DIRECCIÓN REGIONAL DE EDUCACIÓN'}</span><br>
                        <span class="ugel">{escape_html(session.metadata.ugel) or 'UNIDAD DE GESTIÓN EDUCATIVA LOCAL'}</span><br>
                        <span class="agp">ÁREA DE GESTIÓN PEDAGÓGICA</span>
                    </td>
                    <td class="header-logos" style="text-align: right;">
                        {logo_right_html}
                    </td>
                </tr>
            </table>
            
            <div class="divider"></div>
            
            <!-- ════════ TÍTULO PRINCIPAL ════════ -->
            <div class="title-box">
                <h1>SESIÓN DE APRENDIZAJE N° {escape_html(session.metadata.numero_sesion) or '01'}</h1>
                <h2>"{escape_html(session.metadata.titulo) or 'Título de la Sesión'}"</h2>
            </div>
            
            <!-- ════════ DATOS GENERALES ════════ -->
            <table class="content-table">
                <tr>
                    <td class="label-cell">Institución Educativa</td>
                    <td colspan="3">{escape_html(session.metadata.institucion)}</td>
                    <td class="label-cell">Nivel</td>
                    <td>{escape_html(session.metadata.nivel)}</td>
                </tr>
                <tr>
                    <td class="label-cell">Docente</td>
                    <td colspan="3">{escape_html(session.metadata.docente)}</td>
                    <td class="label-cell">Área</td>
                    <td>{escape_html(session.metadata.area)}</td>
                </tr>
                <tr>
                    <td class="label-cell">Grado</td>
                    <td>{escape_html(session.metadata.grado)}</td>
                    <td class="label-cell" style="width: 80px;">Sección</td>
                    <td>{escape_html(session.metadata.seccion)}</td>
                    <td class="label-cell">Unidad / Proyecto</td>
                    <td>{escape_html(session.metadata.unidad)}</td>
                </tr>
                <tr>
                    <td class="label-cell">Fecha</td>
                    <td colspan="3">{escape_html(session.metadata.fecha)}</td>
                    <td class="label-cell">Duración</td>
                    <td>{escape_html(session.metadata.duracion)} min</td>
                </tr>
            </table>
            
            <!-- ════════ I. PROPÓSITO ════════ -->
            <div class="section-title">I. Propósito de la Sesión</div>
            <div class="section-content">
                {escape_html(session.proposito.proposito_texto)}
            </div>
            
            <!-- ════════ II. CONOCIMIENTOS ════════ -->
            <div class="section-title">II. Conocimientos</div>
            <div class="section-content">
                {escape_html(session.proposito.conocimientos)}
            </div>
            
            <!-- ════════ III. PROPÓSITOS DE APRENDIZAJE ════════ -->
            <div class="section-title">III. Propósitos de Aprendizaje</div>
            
            <table class="content-table" style="margin-top: 4px;">
                <tr>
                    <td class="label-cell" style="width: 140px;">Competencia</td>
                    <td><strong>{escape_html(session.proposito.competencia)}</strong></td>
                </tr>
                <tr>
                    <td class="label-cell">Estándar de aprendizaje</td>
                    <td style="font-size: 8.5pt; color: #475569;">{escape_html(session.proposito.estandar)}</td>
                </tr>
            </table>
            
            <table class="content-table">
                <thead>
                    <tr>
                        <th>COMPETENCIAS</th>
                        <th>CAPACIDADES</th>
                        <th>CRITERIOS DE EVALUACIÓN</th>
                        <th>PRODUCTO / EVIDENCIA</th>
                        <th>INSTRUMENTOS</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td style="font-weight: 600;">{escape_html(session.proposito.competencia)}</td>
                        <td><ul class="session-list">{capacidades_html}</ul></td>
                        <td><ul class="session-list">{criterios_html}</ul></td>
                        <td>{escape_html(session.proposito.producto_evidencia)}</td>
                        <td>{escape_html(session.proposito.instrumento)}</td>
                    </tr>
                </tbody>
            </table>
            
            <!-- ════════ COMPETENCIAS TRANSVERSALES ════════ -->
            {"<table class='content-table'><thead><tr><th style='width: 35%'>COMPETENCIAS TRANSVERSALES</th><th>DESEMPEÑOS PRECISADOS / PRODUCTO / INSTRUMENTOS</th></tr></thead><tbody>" + ct_rows_html + "</tbody></table>" if ct_rows_html else ""}
            
            <!-- ════════ ENFOQUES TRANSVERSALES ════════ -->
            {"<table class='content-table'><thead><tr><th style='width: 30%'>ENFOQUES TRANSVERSALES</th><th style='width: 30%'>VALORES</th><th>ACTITUDES OBSERVABLES</th></tr></thead><tbody>" + enfoques_rows_html + "</tbody></table>" if enfoques_rows_html else ""}
            
            <!-- ════════ RECURSOS ════════ -->
            <table class="content-table" style="margin-top: 8px;">
                <tr>
                    <td class="label-cell" style="width: 200px;">Páginas de Texto, otros textos de consulta/Enlaces</td>
                    <td>{escape_html(session.recursos.enlaces)}</td>
                </tr>
                <tr>
                    <td class="label-cell">Materiales y recursos</td>
                    <td>{escape_html(session.recursos.materiales)}</td>
                </tr>
                <tr>
                    <td class="label-cell">Actividades de Refuerzo Escolar</td>
                    <td>{escape_html(session.recursos.refuerzo)}</td>
                </tr>
            </table>
            
            <!-- ════════ IV. SECUENCIA DIDÁCTICA ════════ -->
            <div class="section-title">IV. Secuencia Didáctica (Momentos)</div>
            
            <table class="content-table momentos-table" style="margin-top: 4px;">
                <thead>
                    <tr>
                        <th>MOMENTOS DE LA SESIÓN</th>
                        <th>ESTRATEGIAS / ACTIVIDADES</th>
                        <th>EVALUACIÓN</th>
                    </tr>
                </thead>
                <tbody>
                    <!-- Inicio -->
                    <tr>
                        <td class="momento-label-cell">
                            INICIO
                            <div class="momento-time">TIEMPO: {escape_html(session.momentos.inicio.tiempo_total)} min</div>
                        </td>
                        <td>
                            {inicio_actividades_html}
                        </td>
                        <td class="momento-eval-cell" rowspan="1">
                            EVALUACIÓN FORMATIVA
                        </td>
                    </tr>
                    
                    <!-- Desarrollo (Primer Proceso) -->
                    <tr>
                        <td class="momento-label-cell" rowspan="{cant_procesos}">
                            DESARROLLO
                            <div class="momento-time">TIEMPO: {escape_html(session.momentos.desarrollo.tiempo_total)} min</div>
                        </td>
                        <td>
                            {desarrollo_primero_html}
                        </td>
                        <td class="momento-eval-cell" rowspan="{cant_procesos}">
                            EVALUACIÓN FORMATIVA<br><br>
                            <span style="font-size: 7.5pt; font-weight: normal; color: #64748b;">(Monitoreo activo y retroalimentación)</span>
                        </td>
                    </tr>
                    
                    <!-- Desarrollo (Procesos Siguientes) -->
                    {desarrollo_siguientes_html}
                    
                    <!-- Cierre -->
                    <tr>
                        <td class="momento-label-cell">
                            CIERRE
                            <div class="momento-time">TIEMPO: {escape_html(session.momentos.cierre.tiempo_total)} min</div>
                        </td>
                        <td>
                            {cierre_estrategias_html}
                        </td>
                        <td class="momento-eval-cell">
                            EVALUACIÓN FORMATIVA
                        </td>
                    </tr>
                </tbody>
            </table>
            
            <!-- ════════ FIRMAS ════════ -->
            <div class="firmas-container">
                <div class="firma-box">
                    <div class="firma-linea"></div>
                    <div class="firma-nombre">{escape_html(session.metadata.docente) or 'Docente de la Sesión'}</div>
                    <div class="firma-cargo">Docente de la Sesión</div>
                </div>
                <div class="firma-box">
                    <div class="firma-linea"></div>
                    <div class="firma-nombre">{escape_html(session.metadata.director) or 'Director(a) / Subdirector(a)'}</div>
                    <div class="firma-cargo">Director(a) / Subdirector(a)</div>
                </div>
            </div>
        </div>
        
        <!-- ════════ V. FICHA DE TRABAJO ════════ -->
        {ficha_html}
        
        <!-- ════════ VI. LISTA DE COTEJO ════════ -->
        {lista_cotejo_html}
    </body>
    </html>
    """
    return html_content


@app.post("/exportar-pdf-json")
async def exportar_pdf_json(payload: SesionAprendizajeRequest):
    """
    Genera un archivo PDF a partir del JSON estructurado de la sesión de aprendizaje.
    Intenta utilizar la conversión nativa de Word (docx2pdf) si está disponible en Windows,
    y si falla o no está disponible, cae de vuelta al renderizado con Playwright.
    """
    if payload.token != CONNECTION_TOKEN:
        raise HTTPException(status_code=401, detail="No autorizado: Token de conexión inválido.")

    try:
        titulo = payload.metadata.titulo or "Sesion_de_Aprendizaje"
        filename = re.sub(r'[^a-zA-Z0-9-_\s]', '', titulo).replace(' ', '_')
        nombre_archivo = f"{filename}.pdf"

        # 1. Intentar conversión nativa Word-to-PDF si estamos en Windows
        if sys.platform.startswith('win'):
            try:
                from docx2pdf import convert
                
                # Generamos primero el Word perfecto usando la función premium
                docx_stream = build_docx_from_json(payload)
                
                # Crear archivos temporales
                temp_docx = LOCAL_BIN_DIR / f"temp_{secrets.token_hex(4)}_{filename}.docx"
                temp_pdf = LOCAL_BIN_DIR / f"temp_{secrets.token_hex(4)}_{filename}.pdf"
                
                LOCAL_BIN_DIR.mkdir(exist_ok=True)
                with open(temp_docx, "wb") as f:
                    f.write(docx_stream.read())
                    
                if console:
                    console.print(f"[yellow]⚡ Intentando conversión nativa Word-to-PDF para {nombre_archivo}...[/yellow]")
                
                # Ejecutar la conversión de Word en un hilo separado
                await asyncio.to_thread(convert, str(temp_docx), str(temp_pdf))
                
                # Leer los bytes del PDF resultante
                with open(temp_pdf, "rb") as f:
                    pdf_bytes = f.read()
                    
                # Limpieza de archivos temporales
                try:
                    os.remove(temp_docx)
                    os.remove(temp_pdf)
                except Exception:
                    pass
                    
                if console:
                    console.print(f"[green]✓ [PDF PREMIUM CONVERTIDO] Generado vía Word con éxito: {nombre_archivo}[/green]")
                    
                return Response(
                    content=pdf_bytes,
                    media_type="application/pdf",
                    headers={
                        "Content-Disposition": f"attachment; filename={nombre_archivo}",
                        "Access-Control-Expose-Headers": "Content-Disposition"
                    }
                )
            except Exception as word_err:
                if console:
                    console.print(f"[yellow]⚠️ Falló conversión vía Word: {str(word_err)}. Usando fallback de Chromium...[/yellow]")
                else:
                    print(f"[WARN WORD PDF] Falló conversión: {word_err}. Usando fallback...")

        # 2. Fallback: Renderizado HTML con Playwright (Chromium headless)
        documento_html = build_pdf_html_from_json(payload)

        async with async_playwright() as p:
            ruta_motor = buscar_navegador_compatible()
            launch_args = {"headless": True}
            if ruta_motor:
                launch_args["executable_path"] = ruta_motor
            browser = await p.chromium.launch(**launch_args)
            context = await browser.new_context()
            page = await context.new_page()
            
            await page.set_content(documento_html, wait_until="networkidle")
            
            # Esperar a que MathJax termine de procesar las fórmulas matemáticas (si está presente)
            try:
                await page.evaluate("() => window.MathJax && window.MathJax.startup && window.MathJax.startup.promise")
            except Exception as e:
                print("[WARN MATHJAX WAIT]", str(e))
            
            # Captura a PDF con Playwright aplicando prefer_css_page_size y márgenes en cero (el diseño ya tiene padding)
            pdf_bytes = await page.pdf(
                print_background=True,
                prefer_css_page_size=True,
                margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
                display_header_footer=False
            )
            await browser.close()

        if console:
            console.print(f"[green]✓ [PDF PREMIUM EXPORTADO] Generado vía Chromium con éxito: {nombre_archivo}[/green]")

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={nombre_archivo}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )

    except Exception as e:
        print("[ERROR PDF JSON]", str(e))
        raise HTTPException(status_code=500, detail=f"Fallo al compilar PDF Premium: {str(e)}")


# ────────────────────────────────────────────────────────────────────────
# UTILIDADES PARA CONSTRUCCIÓN DE WORD (.docx) NATIVO
# ────────────────────────────────────────────────────────────────────────

def set_cell_background(cell, hex_color: str):
    """Establece el color de fondo de una celda en Word."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text_direction_vertical(cell):
    """Establece la dirección del texto vertical (abajo a arriba, de izquierda a derecha) en una celda en Word."""
    tcPr = cell._tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), 'btLr')
    tcPr.append(textDirection)


def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Establece márgenes internos (padding) de una celda en dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_table_borders(table, color='CBD5E1', sz='4'):
    """Agrega bordes delgados a toda la tabla. Por defecto gris claro."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        borders.append(border)
    tblPr.append(borders)


def add_table_borders_black(table):
    """Agrega bordes negros medios a toda la tabla (estilo plantilla oficial)."""
    add_table_borders(table, color='000000', sz='8')


def set_run_color(run, hex_color: str):
    """Aplica color hex a un run de texto Word."""
    run.font.color.rgb = RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16)
    )


def set_cell_text_white_bold(cell, text: str, font_size_pt: float = 9):
    """Escribe texto en blanco y negrita en la primera línea de una celda."""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Limpiar párrafo
    for run in p.runs:
        run.text = ''
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def build_docx_from_html(html_content: str) -> io.BytesIO:
    """Parsea recursivamente la estructura HTML y construye un documento .docx nativo."""
    doc = Document()
    
    # Configuración de márgenes estándar
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(30, 41, 59) # Slate-800

    soup = BeautifulSoup(html_content, 'html.parser')

    # Eliminar katex-html para evitar texto duplicado
    for katex_html in soup.find_all(class_='katex-html'):
        katex_html.decompose()

    # Limpiar elementos interactivos del editor que no pertenecen al documento
    for selector in ['no-print', 'add-logo-placeholder', 'btn-remove-logo']:
        for el in soup.find_all(class_=selector):
            el.decompose()

    processed_tags = set()

    def add_runs_to_paragraph(paragraph, element, is_bold=False, is_italic=False):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    if is_bold:
                        run.bold = True
                    if is_italic:
                        run.italic = True
            elif isinstance(child, Tag):
                if child.name == 'br':
                    paragraph.add_run('\n')
                elif child.name in ['strong', 'b']:
                    add_runs_to_paragraph(paragraph, child, is_bold=True, is_italic=is_italic)
                elif child.name in ['em', 'i']:
                    add_runs_to_paragraph(paragraph, child, is_bold=is_bold, is_italic=True)
                elif child.name in ['span', 'a']:
                    add_runs_to_paragraph(paragraph, child, is_bold=is_bold, is_italic=is_italic)
                else:
                    add_runs_to_paragraph(paragraph, child, is_bold=is_bold, is_italic=is_italic)

    def walk_tree(element):
        if element in processed_tags:
            return
        
        if isinstance(element, Tag):
            # TÍTULOS
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                
                run = p.add_run()
                run.bold = True
                run.font.name = 'Arial'
                
                if level == 1:
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(15, 23, 42)
                elif level == 2:
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(30, 41, 59)
                else:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(71, 85, 105)
                
                add_runs_to_paragraph(p, element)
                processed_tags.add(element)
                return

            # PÁRRAFOS
            elif element.name == 'p':
                text_clean = element.get_text(strip=True)
                if not text_clean:
                    processed_tags.add(element)
                    return
                
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                add_runs_to_paragraph(p, element)
                processed_tags.add(element)
                return

            # LISTAS
            elif element.name in ['ul', 'ol']:
                list_style = 'List Bullet' if element.name == 'ul' else 'List Number'
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(style=list_style)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.line_spacing = 1.15
                    add_runs_to_paragraph(p, li)
                processed_tags.add(element)
                return

            # TABLAS — con soporte de rowspan/colspan via merge grid
            elif element.name == 'table':
                html_rows = element.find_all('tr', recursive=True)
                if not html_rows:
                    processed_tags.add(element)
                    return

                # ── Paso 1: Calcular dimensiones reales de la grilla ──
                # Construir una grilla lógica 2D para mapear ocupación de celdas
                num_rows = len(html_rows)
                # Calcular max_cols considerando colspan
                max_cols = 0
                for r in html_rows:
                    total = 0
                    for c in r.find_all(['td', 'th'], recursive=False):
                        cs = int(c.get('colspan', 1) or 1)
                        total += cs
                    if total > max_cols:
                        max_cols = total

                if max_cols == 0:
                    processed_tags.add(element)
                    return

                # Grilla de ocupación: occupied[row][col] = True si ya está ocupada por un rowspan previo
                occupied = [[False] * max_cols for _ in range(num_rows)]

                # ── Paso 2: Crear tabla docx con dimensiones exactas ──
                docx_table = doc.add_table(rows=num_rows, cols=max_cols)
                docx_table.autofit = True
                add_table_borders(docx_table)

                # ── Paso 3: Recorrer filas HTML y mapear celdas con merge ──
                # Almacenar merges pendientes para ejecutar después de llenar contenido
                merges = []  # Lista de (start_row, start_col, end_row, end_col)

                for row_idx, html_row in enumerate(html_rows):
                    html_cells = html_row.find_all(['td', 'th'], recursive=False)
                    col_cursor = 0  # Posición lógica actual en la grilla

                    for html_cell in html_cells:
                        # Avanzar cursor saltando celdas ya ocupadas por rowspan de filas anteriores
                        while col_cursor < max_cols and occupied[row_idx][col_cursor]:
                            col_cursor += 1

                        if col_cursor >= max_cols:
                            break

                        rs = int(html_cell.get('rowspan', 1) or 1)
                        cs = int(html_cell.get('colspan', 1) or 1)

                        # Marcar celdas ocupadas en la grilla
                        for dr in range(rs):
                            for dc in range(cs):
                                target_r = row_idx + dr
                                target_c = col_cursor + dc
                                if target_r < num_rows and target_c < max_cols:
                                    occupied[target_r][target_c] = True

                        # Registrar merge si abarca más de una celda
                        if rs > 1 or cs > 1:
                            end_row = min(row_idx + rs - 1, num_rows - 1)
                            end_col = min(col_cursor + cs - 1, max_cols - 1)
                            merges.append((row_idx, col_cursor, end_row, end_col))

                        # Llenar contenido en la celda de la esquina superior izquierda
                        cell = docx_table.rows[row_idx].cells[col_cursor]
                        is_header = html_cell.name == 'th'

                        bg_color = "F8FAFC" if is_header else "FFFFFF"

                        style_attr = html_cell.get('style', '')
                        class_attr = html_cell.get('class', [])

                        hex_match = re.search(r'background-color:\s*#([A-Fa-f0-9]{6})', style_attr)
                        if hex_match:
                            bg_color = hex_match.group(1).upper()
                        elif 'cell-peru' in class_attr:
                            bg_color = "C0392B"
                        elif 'cell-minedu' in class_attr:
                            bg_color = "2C3E50"

                        set_cell_background(cell, bg_color)
                        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)

                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.1

                        if is_header:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.bold = True
                            if bg_color in ["C0392B", "2C3E50"]:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                run.font.color.rgb = RGBColor(15, 23, 42)
                            add_runs_to_paragraph(p, html_cell)
                        else:
                            add_runs_to_paragraph(p, html_cell)

                        col_cursor += cs

                # ── Paso 4: Ejecutar merges de celdas ──
                for start_r, start_c, end_r, end_c in merges:
                    try:
                        cell_a = docx_table.rows[start_r].cells[start_c]
                        cell_b = docx_table.rows[end_r].cells[end_c]
                        cell_a.merge(cell_b)
                    except Exception:
                        pass  # Merge inválido — ignorar silenciosamente

                doc.add_paragraph().paragraph_format.space_before = Pt(8)
                processed_tags.add(element)
                return

            for child in element.children:
                walk_tree(child)

    walk_tree(soup)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


@app.post("/exportar-docx")
async def exportar_docx(payload: ExportDocxRequest):
    """
    Convierte el HTML de la sesión en un archivo Word (.docx) nativo.
    Requiere token de conexión.
    """
    if payload.token != CONNECTION_TOKEN:
        raise HTTPException(status_code=401, detail="No autorizado: Token de conexión inválido.")

    try:
        # Sanitizar nombre del archivo
        filename = re.sub(r'[^a-zA-Z0-9-_\s]', '', payload.titulo).replace(' ', '_')
        nombre_archivo = f"{filename}.docx"

        # Compilar archivo DOCX
        docx_stream = build_docx_from_html(payload.html_content)

        if console:
            console.print(f"[blue]✓ [WORD EXPORTADO] Generado con éxito: {nombre_archivo}[/blue]")

        return StreamingResponse(
            docx_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={nombre_archivo}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )

    except Exception as e:
        print("[ERROR DOCX]", str(e))
        raise HTTPException(status_code=500, detail=f"Fallo al compilar archivo de Word (.docx): {str(e)}")


def get_image_stream(url: str):
    """Intenta descargar la imagen desde la URL y devuelve un BytesIO. Retorna None si falla o es vacía."""
    if not url or not (url.startswith("http://") or url.startswith("https://")):
        return None
    try:
        req = urllib.request.Request(
            url, 
            headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        )
        with urllib.request.urlopen(req, timeout=3) as response:
            return io.BytesIO(response.read())
    except Exception as e:
        print(f"[WARN LOGO] No se pudo descargar el logo {url}: {e}")
        return None


def format_latex_to_unicode(text: str) -> str:
    """
    Traduce expresiones comunes de LaTeX a caracteres Unicode matemáticos para legibilidad en Word.
    """
    if not text:
        return ""
    
    # Tabla de equivalencias directas de LaTeX a Unicode
    # ── Convertir tablas LaTeX \begin{array}...\end{array} a texto tabular ──
    def convert_latex_table(match):
        content = match.group(1)
        # Obtener filas
        rows_raw = re.split(r'\\\\', content)
        rows = []
        for row_raw in rows_raw:
            row_raw = row_raw.strip()
            if not row_raw:
                continue
            # Separar celdas por &
            cells_raw = row_raw.split('&')
            cells = [c.strip() for c in cells_raw]
            rows.append(cells)
        if not rows:
            return text
        num_cols = max(len(r) for r in rows)
        # Calcular anchos de columna
        col_widths = []
        for c in range(num_cols):
            w = max((len(r[c]) if c < len(r) else 0) for r in rows)
            col_widths.append(max(w, 4))
        # Construir tabla unicode
        lines = []
        sep = '─' * (sum(col_widths) + num_cols * 3 + 1)
        lines.append(sep)
        for r_idx, row in enumerate(rows):
            line_parts = []
            for c_idx in range(num_cols):
                cell_val = row[c_idx] if c_idx < len(row) else ''
                line_parts.append(cell_val.ljust(col_widths[c_idx]))
            lines.append('│ ' + ' │ '.join(line_parts) + ' │')
            if r_idx == 0:
                lines.append(sep)
        lines.append(sep)
        return '\n'.join(lines)

    text = re.sub(r'\\begin\{(?:array|tabular)\}(?:\{[^}]*\})?(.+?)\\end\{(?:array|tabular)\}', convert_latex_table, text, flags=re.DOTALL)

    replacements = {
        r'\pm': '±',
        r'\le': '≤',
        r'\ge': '≥',
        r'\neq': '≠',
        r'\times': '×',
        r'\div': '÷',
        r'\approx': '≈',
        r'\alpha': 'α',
        r'\beta': 'β',
        r'\gamma': 'γ',
        r'\theta': 'θ',
        r'\pi': 'π',
        r'\infty': '∞',
        r'\Delta': 'Δ',
        r'\Sigma': 'Σ',
        r'\rightarrow': '→',
        r'\leftarrow': '←',
        r'\leftrightarrow': '↔',
        r'\Rightarrow': '⇒',
        r'\Leftarrow': '⇐',
        r'\partial': '∂',
        r'\sqrt': '√',
        r'\cdot': '·',
        r'\text{o}': ' o ',
        r'\;': ' ',
        r'\:': ' ',
        r'\,': ' ',
        r'\!': '',
    }
    
    # Remover delimitadores de dólar
    cleaned = text
    cleaned = re.sub(r'\$\$(.*?)\$\$', r'\1', cleaned)
    cleaned = re.sub(r'\$(.*?)\$', r'\1', cleaned)
    
    # Aplicar equivalencias de comandos
    for key, val in replacements.items():
        cleaned = cleaned.replace(key, val)
        
    # Reemplazos de potencias comunes a superíndices unicode
    supers = {'0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴', '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹', 
              '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾', 'n': 'ⁿ', 'x': 'ˣ', 'y': 'ʸ', 'i': 'ⁱ'}
    
    def replace_super(match):
        val = match.group(1)
        return "".join(supers.get(c, c) for c in val)
        
    cleaned = re.sub(r'\^\{?([0-9+\-xyn]+)\}?', replace_super, cleaned)
    
    # Reemplazos de subíndices comunes a unicode
    subs = {'0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄', '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉', 
            '+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎', 'n': 'ₙ', 'x': 'ₓ', 'y': 'ᵧ', 'i': 'ᵢ'}
            
    def replace_sub(match):
        val = match.group(1)
        return "".join(subs.get(c, c) for c in val)
        
    cleaned = re.sub(r'\_\{?([0-9+\-xyni]+)\}?', replace_sub, cleaned)
    
    return cleaned


def preprocess_session_latex(session):
    """
    Recorre el modelo de la sesión y traduce todas las expresiones LaTeX de sus textos
    a caracteres Unicode matemáticos listos para la exportación a Word.
    """
    if not session:
        return
        
    # Metadatos
    if session.metadata:
        session.metadata.titulo = format_latex_to_unicode(session.metadata.titulo)
        session.metadata.institucion = format_latex_to_unicode(session.metadata.institucion)
        session.metadata.docente = format_latex_to_unicode(session.metadata.docente)
        session.metadata.area = format_latex_to_unicode(session.metadata.area)
        session.metadata.unidad = format_latex_to_unicode(session.metadata.unidad)
        
    # Propósitos
    if session.proposito:
        session.proposito.proposito_texto = format_latex_to_unicode(session.proposito.proposito_texto)
        session.proposito.conocimientos = format_latex_to_unicode(session.proposito.conocimientos)
        session.proposito.competencia = format_latex_to_unicode(session.proposito.competencia)
        session.proposito.estandar = format_latex_to_unicode(session.proposito.estandar)
        session.proposito.producto_evidencia = format_latex_to_unicode(session.proposito.producto_evidencia)
        session.proposito.instrumento = format_latex_to_unicode(session.proposito.instrumento)
        
        if session.proposito.capacidades:
            session.proposito.capacidades = [format_latex_to_unicode(c) for c in session.proposito.capacidades]
        if session.proposito.criterios:
            session.proposito.criterios = [format_latex_to_unicode(cr) for cr in session.proposito.criterios]

    # Competencias transversales
    if session.competencias_transversales:
        for ct in session.competencias_transversales:
            ct.titulo = format_latex_to_unicode(ct.titulo)
            if ct.desempenos:
                ct.desempenos = [format_latex_to_unicode(d) for d in ct.desempenos]

    # Enfoques transversales
    if session.enfoques_transversales:
        for et in session.enfoques_transversales:
            et.nombre = format_latex_to_unicode(et.nombre)
            et.valor = format_latex_to_unicode(et.valor)
            et.actitudes = format_latex_to_unicode(et.actitudes)

    # Recursos
    if session.recursos:
        session.recursos.enlaces = format_latex_to_unicode(session.recursos.enlaces)
        session.recursos.materiales = format_latex_to_unicode(session.recursos.materiales)
        session.recursos.refuerzo = format_latex_to_unicode(session.recursos.refuerzo)

    # Momentos
    if getattr(session, 'momentos', None):
        momentos = session.momentos
        # Inicio
        if getattr(momentos, 'inicio', None) and getattr(momentos.inicio, 'actividades', None):
            momentos.inicio.actividades = [format_latex_to_unicode(a) for a in momentos.inicio.actividades]
        # Desarrollo
        if getattr(momentos, 'desarrollo', None) and getattr(momentos.desarrollo, 'procesos', None):
            for proc in momentos.desarrollo.procesos:
                if getattr(proc, 'titulo', None):
                    proc.titulo = format_latex_to_unicode(proc.titulo)
                if getattr(proc, 'contenido', None):
                    proc.contenido = [format_latex_to_unicode(c) for c in proc.contenido]
        # Cierre
        if getattr(momentos, 'cierre', None):
            cierre = momentos.cierre
            if getattr(cierre, 'metacognicion', None):
                cierre.metacognicion = [format_latex_to_unicode(m) for m in cierre.metacognicion]
            if getattr(cierre, 'evaluacion', None):
                cierre.evaluacion = [format_latex_to_unicode(ev) for ev in cierre.evaluacion]
            if getattr(cierre, 'extension', None):
                cierre.extension = [format_latex_to_unicode(ex) for ex in cierre.extension]

    # Ficha de trabajo
    if getattr(session, 'ficha_trabajo', None):
        ficha = session.ficha_trabajo
        if getattr(ficha, 'titulo', None):
            ficha.titulo = format_latex_to_unicode(ficha.titulo)
        if getattr(ficha, 'indicaciones', None):
            ficha.indicaciones = format_latex_to_unicode(ficha.indicaciones)
        if getattr(ficha, 'actividades', None):
            ficha.actividades = format_latex_to_unicode(ficha.actividades)

    # Firmas (si existieran dinámicamente)
    if getattr(session, 'firmas', None):
        for f in session.firmas:
            if getattr(f, 'nombre', None):
                f.nombre = format_latex_to_unicode(f.nombre)
            if getattr(f, 'cargo', None):
                f.cargo = format_latex_to_unicode(f.cargo)


def build_docx_from_json(session: SesionAprendizajeRequest) -> io.BytesIO:
    # Preprocesar LaTeX a Unicode para asegurar compatibilidad matemática y estética en Word
    preprocess_session_latex(session)
    
    doc = Document()
    
    # Margenes exactos A4 y A4 page size
    for s in doc.sections:
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = Inches(0.1)
        s.bottom_margin = Inches(0.1)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # Estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(30, 41, 59)

    # Colores exactos de la plantilla XML
    PEACH = 'F7CAAC'
    BLUE_HDR = 'BDD6EE'
    YELLOW_HDR = 'FFE599'
    GRAY_VAL = 'F2F2F2'
    PEACH_MOM = 'FBE5D5'
    GRAY_MOM = 'EDEDED'
    YELLOW_VAL = 'FFF2CC'
    BULLET_COLORS = ['2980B9', 'C0392B', '27AE60', '8E44AD', '16A085']

    def _label(cell, text):
        set_cell_background(cell, PEACH)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)

    def _val(cell, text):
        set_cell_background(cell, GRAY_VAL)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.1
        p.add_run(text or "").font.size = Pt(9)

    def _hdr(cell, text, bg=BLUE_HDR, sz=9):
        set_cell_background(cell, bg)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(sz)

    def _bullet_cell(cell, items):
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        cell.text = ""
        for i, item in enumerate(items):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            bc = BULLET_COLORS[i % len(BULLET_COLORS)]
            rb = p.add_run(u"\u25cf ")
            rb.font.size = Pt(9)
            rb.font.color.rgb = RGBColor(int(bc[0:2], 16), int(bc[2:4], 16), int(bc[4:6], 16))
            p.add_run(item).font.size = Pt(9)

    def make_vertical_text(text: str) -> str:
        # Stack characters vertically separated by double newlines to match template
        return "\n" + "\n\n".join(list(text)) + "\n"

    def _write_momento_cell(cell, nombre, sub_bullets, tiempo):
        set_cell_background(cell, GRAY_MOM)
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p_main = cell.paragraphs[0]
        rm = p_main.add_run(nombre)
        rm.bold = True
        rm.font.size = Pt(10)
        if sub_bullets:
            p_sub = cell.add_paragraph()
            p_sub.paragraph_format.space_before = Pt(4)
            rs = p_sub.add_run(sub_bullets)
            rs.font.size = Pt(8)
            rs.font.color.rgb = RGBColor(71, 85, 105)
        if tiempo:
            p_t = cell.add_paragraph()
            p_t.paragraph_format.space_before = Pt(6)
            rt = p_t.add_run(u"\u23f1 TIEMPO: " + tiempo + " min")
            rt.bold = True
            rt.font.size = Pt(8.5)
            rt.font.color.rgb = RGBColor(192, 57, 43)

    def _write_vertical_cell(cell, txt):
        set_cell_background(cell, PEACH_MOM)
        set_cell_margins(cell, top=120, bottom=120, left=40, right=40)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(make_vertical_text(txt))
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(120, 60, 20) # Cafe oscuro / oxidado

    # ===============================================================
    # CABECERA INSTITUCIONAL EN EL HEADER NATIVO (TODAS LAS PAGINAS)
    # ===============================================================
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header = section.header

    hdr_tbl = header.add_table(rows=1, cols=7, width=Inches(6.77))
    hdr_tbl.autofit = False
    hdr_tbl.allow_autofit = False

    anchos_hdr  = [Inches(0.85), Inches(0.5), Inches(1.2), Inches(1.15), Inches(1.15), Inches(1.05), Inches(0.87)]
    colores_hdr = [None, 'C00000', '262626', '595959', '7F7F7F', '8FAADC', None]

    for col_idx, hcell in enumerate(hdr_tbl.rows[0].cells):
        hcell.width = anchos_hdr[col_idx]
        tcPr = hcell._tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'right', 'bottom', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcB.append(b)
        tcPr.append(tcB)
        vA = OxmlElement('w:vAlign')
        vA.set(qn('w:val'), 'center')
        tcPr.append(vA)
        if colores_hdr[col_idx]:
            set_cell_background(hcell, colores_hdr[col_idx])

    # Col 0: Logo Peru
    logo_left_stream = get_image_stream(session.metadata.logo_left_url)
    if logo_left_stream:
        try:
            ph = hdr_tbl.cell(0, 0).paragraphs[0]
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ph.add_run().add_picture(logo_left_stream, width=Inches(0.75))
        except Exception:
            pass

    # Col 1: PERU
    set_cell_text_white_bold(hdr_tbl.cell(0, 1), "PERU", 7)

    # Col 2: MINEDU
    ph2 = hdr_tbl.cell(0, 2).paragraphs[0]
    ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = ph2.add_run("MINISTERIO DE EDUCACION")
    r2.bold = True
    r2.font.size = Pt(7.5)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Col 3: DRE
    ph3 = hdr_tbl.cell(0, 3).paragraphs[0]
    ph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = ph3.add_run(session.metadata.dre or "DIRECCION REGIONAL DE EDUCACION")
    r3.bold = True
    r3.font.size = Pt(7)
    r3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Col 4: UGEL
    ph4 = hdr_tbl.cell(0, 4).paragraphs[0]
    ph4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = ph4.add_run(session.metadata.ugel or "UNIDAD DE GESTION EDUCATIVA LOCAL")
    r4.bold = True
    r4.font.size = Pt(7)
    r4.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Col 5: AGP
    ph5 = hdr_tbl.cell(0, 5).paragraphs[0]
    ph5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = ph5.add_run("AREA DE GESTION PEDAGOGICA")
    r5.bold = True
    r5.font.size = Pt(7)
    r5.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Col 6: Logo Regional
    logo_right_stream = get_image_stream(session.metadata.logo_regional_url)
    if logo_right_stream:
        try:
            ph6 = hdr_tbl.cell(0, 6).paragraphs[0]
            ph6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ph6.add_run().add_picture(logo_right_stream, width=Inches(0.75))
        except Exception:
            pass

    # Limpiar parrafo nativo del header
    if header.paragraphs:
        hp = header.paragraphs[0]
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)

    # ===============================================================
    # BANNER: Sesion N + Titulo + Proposito + Conocimientos
    # ===============================================================
    banner = doc.add_table(rows=4, cols=1)
    banner.autofit = True
    add_table_borders_black(banner)

    # Fila 0: SESION DE APRENDIZAJE
    cn = banner.cell(0, 0)
    set_cell_background(cn, 'D9E1F2')
    set_cell_margins(cn, top=80, bottom=80, left=120, right=120)
    pn = cn.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = pn.add_run("SESION DE APRENDIZAJE N° " + (session.metadata.numero_sesion or '01'))
    rn.bold = True
    rn.font.size = Pt(12)

    # Fila 1: TITULO encabezado
    ct1 = banner.cell(1, 0)
    set_cell_background(ct1, BLUE_HDR)
    set_cell_margins(ct1, top=60, bottom=60, left=120, right=120)
    pt1 = ct1.paragraphs[0]
    pt1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt1 = pt1.add_run("TITULO DE LA SESION")
    rt1.bold = True
    rt1.font.size = Pt(9.5)

    # Fila 2: valor titulo
    ct2 = banner.cell(2, 0)
    set_cell_background(ct2, 'FFFFFF')
    set_cell_margins(ct2, top=80, bottom=80, left=140, right=140)
    pt2 = ct2.paragraphs[0]
    pt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt2 = pt2.add_run('"' + (session.metadata.titulo or 'Titulo de la Sesion') + '"')
    rt2.bold = True
    rt2.italic = True
    rt2.font.size = Pt(11)

    # Fila 3: tabla interna 2 columnas (Proposito | Conocimientos)
    ct3 = banner.cell(3, 0)
    set_cell_margins(ct3, top=0, bottom=0, left=0, right=0)
    inner = ct3.add_table(rows=2, cols=2)
    inner.autofit = False
    inner.allow_autofit = False

    ci_ph = inner.cell(0, 0)
    ci_ph.width = Inches(4.4)
    _hdr(ci_ph, "PROPOSITO DE LA SESION:", bg=BLUE_HDR, sz=9)

    ci_ch = inner.cell(0, 1)
    ci_ch.width = Inches(2.37)
    _hdr(ci_ch, "CONOCIMIENTOS:", bg=PEACH, sz=9)

    ci_pv = inner.cell(1, 0)
    ci_pv.width = Inches(4.4)
    set_cell_background(ci_pv, 'FFFFFF')
    set_cell_margins(ci_pv, top=80, bottom=80, left=140, right=140)
    ci_pv.paragraphs[0].add_run(session.proposito.proposito_texto or "")

    ci_cv = inner.cell(1, 1)
    ci_cv.width = Inches(2.37)
    set_cell_background(ci_cv, 'FFFFFF')
    set_cell_margins(ci_cv, top=80, bottom=80, left=140, right=140)
    ci_cv.paragraphs[0].add_run(session.proposito.conocimientos or "")

    # ===============================================================
    # I. DATOS GENERALES
    # ===============================================================
    doc.add_paragraph().paragraph_format.space_before = Pt(10)
    h1 = doc.add_paragraph()
    h1.paragraph_format.keep_with_next = True
    r1 = h1.add_run("I. DATOS GENERALES")
    r1.bold = True
    r1.font.size = Pt(11)

    dg = doc.add_table(rows=3, cols=13)
    dg.autofit = False
    dg.allow_autofit = False
    add_table_borders_black(dg)

    # Row 0 cells: label(sp=3 IE), val(sp=4), label(sp=2 Nivel), val(sp=4)
    c_ie_lbl = dg.cell(0, 0).merge(dg.cell(0, 2))
    c_ie_val = dg.cell(0, 3).merge(dg.cell(0, 6))
    c_niv_lbl = dg.cell(0, 7).merge(dg.cell(0, 8))
    c_niv_val = dg.cell(0, 9).merge(dg.cell(0, 12))
    
    _label(c_ie_lbl, "Institución Educativa")
    _val(c_ie_val, session.metadata.institucion)
    _label(c_niv_lbl, "Nivel")
    _val(c_niv_val, session.metadata.nivel)

    # Row 1 cells: label(sp=1 Docente), val(sp=6), label(sp=2 Area), val(sp=2), label(sp=1 Unidad), val(sp=1)
    c_doc_lbl = dg.cell(1, 0)
    c_doc_val = dg.cell(1, 1).merge(dg.cell(1, 6))
    c_are_lbl = dg.cell(1, 7).merge(dg.cell(1, 8))
    c_are_val = dg.cell(1, 9).merge(dg.cell(1, 10))
    c_uni_lbl = dg.cell(1, 11)
    c_uni_val = dg.cell(1, 12)

    _label(c_doc_lbl, "Docente")
    _val(c_doc_val, session.metadata.docente)
    _label(c_are_lbl, "Área")
    _val(c_are_val, session.metadata.area)
    _label(c_uni_lbl, "Unidad/ Proyecto")
    _val(c_uni_val, session.metadata.unidad)

    # Row 2 cells: label(sp=1 Grado), val(sp=1), label(sp=2 Seccion), val(sp=1), label(sp=1 Fecha), val(sp=2), label(sp=2 Duracion), val(sp=3)
    c_gra_lbl = dg.cell(2, 0)
    c_gra_val = dg.cell(2, 1)
    c_sec_lbl = dg.cell(2, 2).merge(dg.cell(2, 3))
    c_sec_val = dg.cell(2, 4)
    c_fec_lbl = dg.cell(2, 5)
    c_fec_val = dg.cell(2, 6).merge(dg.cell(2, 7))
    c_dur_lbl = dg.cell(2, 8).merge(dg.cell(2, 9))
    c_dur_val = dg.cell(2, 10).merge(dg.cell(2, 12))

    _label(c_gra_lbl, "Grado")
    _val(c_gra_val, session.metadata.grado)
    _label(c_sec_lbl, "Sección")
    _val(c_sec_val, session.metadata.seccion)
    _label(c_fec_lbl, "Fecha")
    _val(c_fec_val, session.metadata.fecha)
    _label(c_dur_lbl, "Duración")
    _val(c_dur_val, (session.metadata.duracion + " min") if session.metadata.duracion else "")

    # Set exact column widths for the 13 columns of the general data table
    anchos_dg = [Inches(0.6), Inches(0.4), Inches(0.4), Inches(0.55), Inches(0.55), Inches(0.55), Inches(0.55), Inches(0.45), Inches(0.45), Inches(0.55), Inches(0.55), Inches(0.55), Inches(0.52)]
    for row in dg.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = anchos_dg[ci]

    # ===============================================================
    # II. PROPOSITOS DE APRENDIZAJE
    # ===============================================================
    doc.add_paragraph().paragraph_format.space_before = Pt(10)
    h2 = doc.add_paragraph()
    h2.paragraph_format.keep_with_next = True
    r2pa = h2.add_run("II. PROPOSITOS DE APRENDIZAJE")
    r2pa.bold = True
    r2pa.font.size = Pt(11)

    # Subtabla Competencia / Estandar
    ce_tbl = doc.add_table(rows=2, cols=2)
    ce_tbl.autofit = True
    add_table_borders_black(ce_tbl)
    _label(ce_tbl.cell(0, 0), "Competencia")
    ce_tbl.cell(0, 0).width = Inches(1.8)
    _val(ce_tbl.cell(0, 1), session.proposito.competencia)
    _label(ce_tbl.cell(1, 0), "Estandar de aprendizaje")
    _val(ce_tbl.cell(1, 1), session.proposito.estandar)

    # Matriz de Propositos
    doc.add_paragraph().paragraph_format.space_before = Pt(6)
    mx = doc.add_table(rows=2, cols=5)
    mx.autofit = False
    mx.allow_autofit = False
    add_table_borders_black(mx)

    for i, ht in enumerate(["COMPETENCIAS", "CAPACIDADES", "CRITERIOS DE EVALUACION", "PRODUCTO / EVIDENCIA", "INSTRUMENTOS DE EVALUACION"]):
        _hdr(mx.cell(0, i), ht, bg=BLUE_HDR, sz=8.5)

    _val(mx.cell(1, 0), session.proposito.competencia)
    if mx.cell(1, 0).paragraphs[0].runs:
        mx.cell(1, 0).paragraphs[0].runs[0].bold = True

    _bullet_cell(mx.cell(1, 1), session.proposito.capacidades)
    _bullet_cell(mx.cell(1, 2), session.proposito.criterios)
    _val(mx.cell(1, 3), session.proposito.producto_evidencia)
    _val(mx.cell(1, 4), session.proposito.instrumento)

    for row in mx.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = [Inches(1.35), Inches(1.4), Inches(1.6), Inches(1.3), Inches(1.12)][ci]
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15

    # Competencias Transversales
    if session.competencias_transversales:
        doc.add_paragraph().paragraph_format.space_before = Pt(8)
        ct_tbl = doc.add_table(rows=1 + len(session.competencias_transversales), cols=2)
        ct_tbl.autofit = False
        ct_tbl.allow_autofit = False
        add_table_borders_black(ct_tbl)
        _hdr(ct_tbl.cell(0, 0), "COMPETENCIAS TRANSVERSALES", bg=BLUE_HDR, sz=9)
        _hdr(ct_tbl.cell(0, 1), "DESEMPENOS PRECISADOS / PRODUCTO / INSTRUMENTOS", bg=BLUE_HDR, sz=9)
        for i, ct in enumerate(session.competencias_transversales):
            cell_t = ct_tbl.cell(i + 1, 0)
            set_cell_margins(cell_t, top=80, bottom=80, left=120, right=120)
            cell_t.text = ""
            rct = cell_t.add_paragraph().add_run(ct.titulo)
            rct.bold = True
            rct.font.size = Pt(9)
            _bullet_cell(ct_tbl.cell(i + 1, 1), ct.desempenos)
        for row in ct_tbl.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = [Inches(2.4), Inches(4.37)][ci]

    # Enfoques Transversales
    if session.enfoques_transversales:
        doc.add_paragraph().paragraph_format.space_before = Pt(8)
        et_tbl = doc.add_table(rows=1 + len(session.enfoques_transversales), cols=3)
        et_tbl.autofit = False
        et_tbl.allow_autofit = False
        add_table_borders_black(et_tbl)
        for i, ht in enumerate(["ENFOQUES TRANSVERSALES", "VALORES", "ACTITUDES O ACCIONES OBSERVABLES"]):
            _hdr(et_tbl.cell(0, i), ht, bg=BLUE_HDR, sz=9)
        for i, enf in enumerate(session.enfoques_transversales):
            c0 = et_tbl.cell(i + 1, 0)
            set_cell_margins(c0, top=80, bottom=80, left=120, right=120)
            c0.text = ""
            re0 = c0.add_paragraph().add_run(enf.nombre)
            re0.bold = True
            re0.font.size = Pt(9)
            for ci2, val in enumerate([enf.valor, enf.actitudes]):
                cv = et_tbl.cell(i + 1, ci2 + 1)
                set_cell_margins(cv, top=80, bottom=80, left=120, right=120)
                cv.text = ""
                cv.add_paragraph().add_run(val or "").font.size = Pt(9)
        for row in et_tbl.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = [Inches(1.9), Inches(1.7), Inches(3.17)][ci]

    # Recursos y Materiales
    doc.add_paragraph().paragraph_format.space_before = Pt(8)
    rec = doc.add_table(rows=3, cols=2)
    rec.autofit = True
    add_table_borders_black(rec)
    for i, (lbl, val) in enumerate([
        ("Paginas de Texto, otros textos / Enlace web, etc.", session.recursos.enlaces),
        ("Materiales y recursos", session.recursos.materiales),
        ("Actividades de Refuerzo Escolar (N ficha y Titulo)", session.recursos.refuerzo)
    ]):
        _label(rec.cell(i, 0), lbl)
        rec.cell(i, 0).width = Inches(3.0)
        _val(rec.cell(i, 1), val or "")

    # ===============================================================
    # III. SECUENCIA DIDACTICA (MOMENTOS) - 4 COLUMNAS PREMIUM
    # ===============================================================
    doc.add_paragraph().paragraph_format.space_before = Pt(12)
    h3 = doc.add_paragraph()
    h3.paragraph_format.keep_with_next = True
    r3 = h3.add_run("III. SECUENCIA DIDACTICA (MOMENTOS DE LA SESION)")
    r3.bold = True
    r3.font.size = Pt(11)

    procs = session.momentos.desarrollo.procesos
    n_proc = len(procs) if procs else 1
    n_rows = 1 + 1 + n_proc + 1  # header + inicio + desarrollo(s) + cierre

    # Tabla de 4 columnas: Momentos | Motivacion | Contenido | Evaluacion
    mt = doc.add_table(rows=n_rows, cols=4)
    mt.autofit = False
    mt.allow_autofit = False
    add_table_borders_black(mt)

    # Cabeceras
    _hdr(mt.cell(0, 0), "MOMENTOS DE LA SESION", bg=BLUE_HDR, sz=9.5)
    
    # Combinar columnas 1, 2 y 3 para la cabecera "ESTRATEGIAS / ACTIVIDADES"
    cell_est_hdr = mt.cell(0, 1).merge(mt.cell(0, 3))
    _hdr(cell_est_hdr, "ESTRATEGIAS / ACTIVIDADES", bg=BLUE_HDR, sz=9.5)

    # Fila 1: INICIO
    _write_momento_cell(mt.cell(1, 0), "INICIO",
        "Saberes Previos\nProblematizacion\nMotivacion",
        session.momentos.inicio.tiempo_total or "")
    _write_vertical_cell(mt.cell(1, 1), "MOTIVACION")
    _write_vertical_cell(mt.cell(1, 3), "EVALUACION")
    
    cell_ini = mt.cell(1, 2)
    set_cell_margins(cell_ini, top=120, bottom=120, left=140, right=140)
    cell_ini.text = ""
    for act in session.momentos.inicio.actividades:
        p = cell_ini.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(act).font.size = Pt(9.5)

    # Filas 2 a 2+n_proc-1: DESARROLLO
    _write_momento_cell(mt.cell(2, 0), "DESARROLLO",
        "Gestion y Acompanamiento del Desarrollo de las Competencias\n(Procesos didacticos del Area)",
        session.momentos.desarrollo.tiempo_total or "")
    _write_vertical_cell(mt.cell(2, 1), "MOTIVACION")
    _write_vertical_cell(mt.cell(2, 3), "EVALUACION")

    for idx in range(n_proc):
        rn = 2 + idx
        cell_des = mt.cell(rn, 2)
        set_cell_margins(cell_des, top=120, bottom=120, left=140, right=140)
        cell_des.text = ""
        if procs:
            proc = procs[idx]
            ppt = cell_des.add_paragraph()
            ppt.paragraph_format.space_after = Pt(4)
            rpt = ppt.add_run(proc.titulo.upper())
            rpt.bold = True
            rpt.font.size = Pt(9.5)
            rpt.font.color.rgb = RGBColor(192, 57, 43)
            for par in proc.contenido:
                ppr = cell_des.add_paragraph()
                ppr.paragraph_format.space_after = Pt(4)
                ppr.paragraph_format.line_spacing = 1.15
                ppr.add_run(par).font.size = Pt(9.5)
        else:
            cell_des.add_paragraph().add_run("Gestion y Acompanamiento del Desarrollo de Competencias...").font.size = Pt(9.5)

        # Para las filas adicionales, darles formato de fondo en las celdas de Motivacion y Evaluacion
        if idx > 0:
            _write_vertical_cell(mt.cell(rn, 1), "MOTIVACION")
            _write_vertical_cell(mt.cell(rn, 3), "EVALUACION")

    if n_proc > 1:
        # Fusionar verticalmente columnas 0, 1 y 3 en la seccion Desarrollo
        mt.cell(2, 0).merge(mt.cell(2 + n_proc - 1, 0))
        mt.cell(2, 1).merge(mt.cell(2 + n_proc - 1, 1))
        mt.cell(2, 3).merge(mt.cell(2 + n_proc - 1, 3))

    # Fila CIERRE
    rc = n_rows - 1
    _write_momento_cell(mt.cell(rc, 0), "CIERRE",
        "Evaluacion (Reflexion sobre lo aprendido)\nAcciones de reforzamiento o indagacion",
        session.momentos.cierre.tiempo_total or "")
    _write_vertical_cell(mt.cell(rc, 1), "MOTIVACION")
    _write_vertical_cell(mt.cell(rc, 3), "EVALUACION")

    cell_cie = mt.cell(rc, 2)
    set_cell_margins(cell_cie, top=120, bottom=120, left=140, right=140)
    cell_cie.text = ""

    cierre = session.momentos.cierre
    if cierre.metacognicion:
        plbl = cell_cie.add_paragraph()
        rlbl = plbl.add_run("Metacognicion:")
        rlbl.bold = True
        rlbl.font.size = Pt(9.5)
        for m in cierre.metacognicion:
            pi = cell_cie.add_paragraph(style='List Bullet')
            pi.paragraph_format.space_after = Pt(2)
            pi.add_run(m).font.size = Pt(9.5)
    if cierre.evaluacion:
        plbl = cell_cie.add_paragraph()
        plbl.paragraph_format.space_before = Pt(4)
        rlbl = plbl.add_run("Evaluacion formativa:")
        rlbl.bold = True
        rlbl.font.size = Pt(9.5)
        for ev in cierre.evaluacion:
            pi = cell_cie.add_paragraph(style='List Bullet')
            pi.paragraph_format.space_after = Pt(2)
            pi.add_run(ev).font.size = Pt(9.5)
    if cierre.extension:
        plbl = cell_cie.add_paragraph()
        plbl.paragraph_format.space_before = Pt(4)
        rlbl = plbl.add_run("Extension para casa:")
        rlbl.bold = True
        rlbl.font.size = Pt(9.5)
        for ext in cierre.extension:
            pi = cell_cie.add_paragraph(style='List Bullet')
            pi.paragraph_format.space_after = Pt(2)
            pi.add_run(ext).font.size = Pt(9.5)

    # Set exact width for 4 columns of Moments table
    anchos_mom = [Inches(1.2), Inches(0.35), Inches(4.87), Inches(0.35)]
    for row in mt.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = anchos_mom[ci]

    # ===============================================================
    # FIRMAS DE LA SESION
    # ===============================================================
    doc.add_paragraph().paragraph_format.space_before = Pt(40)
    ft = doc.add_table(rows=1, cols=2)
    ft.autofit = True
    for cell in ft.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcB.append(b)
        tcPr.append(tcB)
    pfd = ft.cell(0, 0).paragraphs[0]
    pfd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pfd.add_run("_______________________________\n").bold = True
    rnd = pfd.add_run((session.metadata.docente or "Docente de la Sesion") + "\n")
    rnd.bold = True
    rnd.font.size = Pt(9.5)
    rcd = pfd.add_run("Docente de la Sesion")
    rcd.font.size = Pt(8.5)
    rcd.font.color.rgb = RGBColor(100, 116, 139)

    pfdi = ft.cell(0, 1).paragraphs[0]
    pfdi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pfdi.add_run("_______________________________\n").bold = True
    rndi = pfdi.add_run((session.metadata.director or "Director(a) / Subdirector(a)") + "\n")
    rndi.bold = True
    rndi.font.size = Pt(9.5)
    rcdi = pfdi.add_run("Director(a) / Subdirector(a)")
    rcdi.font.size = Pt(8.5)
    rcdi.font.color.rgb = RGBColor(100, 116, 139)

    # ===============================================================
    # FICHA DE TRABAJO (SI EXISTE)
    # ===============================================================
    if session.ficha_trabajo:
        doc.add_page_break()
        pftb = doc.add_paragraph()
        pftb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pftb.paragraph_format.space_before = Pt(20)
        pftb.paragraph_format.space_after  = Pt(12)
        rftt = pftb.add_run("FICHA DE TRABAJO INDEPENDIENTE PARA EL ESTUDIANTE")
        rftt.bold = True
        rftt.font.size = Pt(12)
        ftt = doc.add_table(rows=1, cols=2)
        ftt.autofit = True
        for cell in ftt.rows[0].cells:
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bb = OxmlElement('w:bottom')
            bb.set(qn('w:val'), 'single')
            bb.set(qn('w:sz'), '8')
            bb.set(qn('w:color'), '3498DB')
            tcBorders.append(bb)
            tcPr.append(tcBorders)
        ftt.cell(0, 0).text = "Nombre: __________________________________________________"
        ftt.cell(0, 0).paragraphs[0].runs[0].bold = True
        ftt.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(10)
        ftt.cell(0, 1).text = "Grado y Seccion: ________________"
        ftt.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ftt.cell(0, 1).paragraphs[0].runs[0].bold = True
        ftt.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(10)

        doc.add_paragraph().paragraph_format.space_before = Pt(14)
        pfta = doc.add_paragraph()
        rfta = pfta.add_run("Actividad: " + (session.ficha_trabajo.titulo or 'Mi Ficha Practica'))
        rfta.bold = True
        rfta.font.size = Pt(12)
        rfta.font.color.rgb = RGBColor(41, 128, 185)
        pfti = doc.add_paragraph()
        pfti.paragraph_format.space_before = Pt(6)
        pfti.paragraph_format.space_after  = Pt(14)
        rftil = pfti.add_run("Indicaciones: ")
        rftil.bold = True
        rftil.font.size = Pt(9.5)
        rftiv = pfti.add_run(session.ficha_trabajo.indicaciones or "Realiza la actividad segun las indicaciones.")
        rftiv.italic = True
        rftiv.font.size = Pt(9.5)
        pftc = doc.add_paragraph()
        pftc.paragraph_format.left_indent = Inches(0.1)
        ctxt = re.sub(r'<[^>]*>', '', session.ficha_trabajo.actividades or "")
        pftc.add_run(ctxt)

    # ===============================================================
    # LISTA DE COTEJO (INSTRUMENTO DE EVALUACION)
    # ===============================================================
    from docx.enum.table import WD_TABLE_ALIGNMENT
    doc.add_page_break()

    plct = doc.add_paragraph()
    plct.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rlct = plct.add_run("INSTRUMENTO DE EVALUACION")
    rlct.bold = True
    rlct.font.size = Pt(12)

    plcs = doc.add_paragraph()
    plcs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    plcs.paragraph_format.space_after = Pt(14)
    rlcs = plcs.add_run("LISTA DE COTEJO - " + (session.metadata.grado or '2 grado') + " " + (session.metadata.seccion or 'A'))
    rlcs.bold = True
    rlcs.font.size = Pt(10.5)
    rlcs.font.color.rgb = RGBColor(100, 116, 139)

    alumnos = session.alumnos if session.alumnos else [f"Estudiante {i+1}" for i in range(30)]
    criterios = session.proposito.criterios if session.proposito.criterios else [
        "Expresa con diversas representaciones la comprension sobre el tema.",
        "Ordena y organiza conceptos clave para resolver problemas.",
        "Emplea estrategias y procedimientos diversos para realizar las tareas.",
        "Halla y valida soluciones utilizando criterios y conocimientos del area."
    ]

    CRIT_COLORS   = ['D9E1F2', 'FADBD8', 'D5F5E3', 'FCF3CF', 'FDE8D8', 'E8DAEF']
    SUBCRIT_COLORS = ['BDD7EE', 'FADBD8', 'A9DFBF', 'F9E79F', 'FAD7A0', 'D7BDE2']

    num_cols = 2 + len(criterios) * 2
    lct = doc.add_table(rows=2 + len(alumnos), cols=num_cols)
    lct.alignment = WD_TABLE_ALIGNMENT.CENTER
    lct.autofit = False

    lct.cell(0, 0).merge(lct.cell(1, 0))
    lct.cell(0, 1).merge(lct.cell(1, 1))
    lct.cell(0, 0).text = "N"
    lct.cell(0, 1).text = "ESTUDIANTES"
    for ci, crit in enumerate(criterios):
        sc = 2 + ci * 2
        lct.cell(0, sc).merge(lct.cell(0, sc + 1))
        lct.cell(0, sc).text = crit
        lct.cell(1, sc).text = "SI"
        lct.cell(1, sc + 1).text = "NO"

    def _lcfmt(cell, w, fs, bold=False, ctr=True, bg=None):
        cell.width = Inches(w)
        tcPr = cell._tc.get_or_add_tcPr()
        vA = OxmlElement('w:vAlign')
        vA.set(qn('w:val'), 'center')
        tcPr.append(vA)
        tcB = OxmlElement('w:tcBorders')
        for bn in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{bn}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
            tcB.append(b)
        tcPr.append(tcB)
        if bg:
            set_cell_background(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ctr else WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0] if p.runs else p.add_run(cell.text)
        if not p.runs:
            cell.text = ""
        run.font.name = 'Arial'
        run.font.size = Pt(fs)
        run.bold = bold

    _lcfmt(lct.cell(0, 0), 0.35, 8.5, bold=True, ctr=True,  bg='FFF2CC')
    _lcfmt(lct.cell(0, 1), 2.2,  8.5, bold=True, ctr=False, bg='FFF2CC')
    for ci, crit in enumerate(criterios):
        sc = 2 + ci * 2
        _lcfmt(lct.cell(0, sc),     0.7,  7.5, bold=True, ctr=True, bg=CRIT_COLORS[ci % len(CRIT_COLORS)])
        _lcfmt(lct.cell(1, sc),     0.35, 8,   bold=True, ctr=True, bg=SUBCRIT_COLORS[ci % len(SUBCRIT_COLORS)])
        _lcfmt(lct.cell(1, sc + 1), 0.35, 8,   bold=True, ctr=True, bg=SUBCRIT_COLORS[ci % len(SUBCRIT_COLORS)])
    for ri, stud in enumerate(alumnos):
        rn = 2 + ri
        lct.cell(rn, 0).text = str(ri + 1)
        lct.cell(rn, 1).text = "" if stud.startswith("Estudiante ") else stud
        _lcfmt(lct.cell(rn, 0), 0.35, 8.5, bold=True,  ctr=True)
        _lcfmt(lct.cell(rn, 1), 2.2,  8.5, bold=False, ctr=False)
        for ci in range(len(criterios)):
            sc = 2 + ci * 2
            _lcfmt(lct.cell(rn, sc),     0.35, 8, ctr=True)
            _lcfmt(lct.cell(rn, sc + 1), 0.35, 8, ctr=True)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

@app.post("/exportar-docx-json")
async def exportar_docx_json(payload: SesionAprendizajeRequest):
    """
    Exporta una sesión de aprendizaje completa desde JSON a un archivo Word (.docx) nativo premium.
    Requiere token de conexión.
    """
    if payload.token != CONNECTION_TOKEN:
        raise HTTPException(status_code=401, detail="No autorizado: Token de conexión inválido.")

    try:
        titulo = payload.metadata.titulo or "Sesion_de_Aprendizaje"
        filename = re.sub(r'[^a-zA-Z0-9-_\s]', '', titulo).replace(' ', '_')
        nombre_archivo = f"{filename}.docx"

        # Compilar archivo DOCX desde JSON nativo
        docx_stream = build_docx_from_json(payload)

        if console:
            console.print(f"[green]✓ [WORD PREMIUM EXPORTADO] Generado nativamente con éxito: {nombre_archivo}[/green]")

        return StreamingResponse(
            docx_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={nombre_archivo}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )

    except Exception as e:
        print("[ERROR DOCX JSON]", str(e))
        raise HTTPException(status_code=500, detail=f"Fallo al compilar archivo de Word (.docx) nativo: {str(e)}")


# ────────────────────────────────────────────────────────────────────────
# INICIALIZACIÓN, COMPROBACIÓN DE CHROMIUM Y ARRANQUE
# ────────────────────────────────────────────────────────────────────────

def buscar_navegador_compatible():
    """Busca un ejecutable de Chromium en la carpeta local o en el sistema."""
    # 1. Prioridad Máxima: Verificar si ya existe en nuestra carpeta portable './bin'
    if CHROMIUM_EXE.exists():
        return str(CHROMIUM_EXE)

    # 2. Prioridad Secundaria: Buscar navegadores instalados en Windows
    if sys.platform.startswith('win'):
        rutas_sistema = [
            r"C:\Program Files\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files\BraveSoftware\Brave-Browser\Application\brave.exe",
            os.path.expandvars(r"%LOCALAPPDATA%\Google\Chrome\Application\chrome.exe"),
            r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
        ]
        for ruta in rutas_sistema:
            if Path(ruta).exists():
                return ruta
    return None


def descargar_chromium_nativo():
    """Descarga Chromium usando urllib para no congelar el .exe y muestra barra de progreso."""
    LOCAL_BIN_DIR.mkdir(exist_ok=True)
    zip_path = LOCAL_BIN_DIR / "chromium.zip"
    
    # URL directa de Google APIs (Versión ligera y estable para Windows x64)
    url_chromium = "https://storage.googleapis.com/chromium-browser-snapshots/Win_x64/1182249/chrome-win.zip"

    print("\n⚠️  [MOTOR INCOMPLETO] No se detectó ningún navegador compatible (Chrome/Brave/Edge) ni motor local.")
    print("Iniciando descarga de motor Chromium portable (aprox. 140MB)...")
    
    last_percent = -1
    def progreso_download(block_num, block_size, total_size):
        nonlocal last_percent
        if total_size > 0:
            completed = block_num * block_size
            percent = int((completed / total_size) * 100)
            if percent != last_percent and percent % 5 == 0:  # Cada 5%
                last_percent = percent
                completed_mb = completed / (1024 * 1024)
                total_mb = total_size / (1024 * 1024)
                print(f"Descargando Chromium: {percent}% completado ({completed_mb:.1f} MB de {total_mb:.1f} MB)")

    # Extracción del ZIP de forma nativa
    try:
        urllib.request.urlretrieve(url_chromium, zip_path, reporthook=progreso_download)
        print("✓ Descarga completada. Extrayendo motor portable...")
        
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(LOCAL_BIN_DIR)
            
        # Eliminar el ZIP basura para ahorrar espacio
        if zip_path.exists():
            os.remove(zip_path)
            
        print("✓ Motor Chromium extraído y listo para usar en ./bin/chrome-win/\n")
    except Exception as e:
        print(f"❌ Error al descargar o extraer Chromium: {str(e)}")


@app.on_event("startup")
async def startup_event():
    """Evento que se dispara al iniciar FastAPI para mostrar la URL de conexión segura en los logs."""
    target_url = f"https://sesiones.sypablitodp.site/conexion.html?token={CONNECTION_TOKEN}"
    print(f"🌐 [MOTOR ONLINE] Servidor de exportación corriendo en http://localhost:8000")
    print(f"🔗 [ENLACE SEGURO] URL de vinculación segura:\n{target_url}\n")


class TerminalRedirector:
    def __init__(self, text_widget):
        self.text_widget = text_widget
        self.ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')

    def write(self, string):
        clean_string = self.ansi_escape.sub('', string)
        if not clean_string: return
        
        try:
            self.text_widget.configure(state='normal')
            
            # Limitar el historial de la terminal a las últimas 300 líneas para evitar fugas de memoria
            try:
                line_count = int(self.text_widget.index('end-1c').split('.')[0])
                if line_count > 300:
                    self.text_widget.delete('1.0', '100.0') # Borra las primeras 100 líneas antiguas
            except Exception:
                pass

            # Pintar de verde si detecta éxito, rojo si es error, blanco el resto
            tag = "muted"
            if "✓" in clean_string or "ONLINE" in clean_string or "Conectado" in clean_string or "enlazado" in clean_string: tag = "green"
            if "❌" in clean_string or "ERROR" in clean_string or "Fallo" in clean_string: tag = "red"
            
            self.text_widget.insert('end', clean_string, tag)
            self.text_widget.see('end')
            self.text_widget.configure(state='disabled')
        except Exception:
            pass

    def flush(self): pass


def start_gui():
    """Interfaz gráfica ultra-minimalista tipo terminal hacker."""
    root = tk.Tk()
    root.title("S.Y. PABLITO_DP - Servidor Local")
    root.geometry("850x500")
    root.configure(bg="#050505") # Negro profundo
    root.resizable(False, False)

    # Cargar icono si existe
    ico_path = BASE_DIR / "assets" / "logo.ico"
    if not ico_path.exists():
        ico_path = EXE_DIR.parent / "assets" / "logo.ico"
    if ico_path.exists():
        try:
            root.iconbitmap(str(ico_path))
        except Exception:
            pass

    # Interceptar el evento de cierre de ventana para mostrar advertencia
    from tkinter import messagebox
    def on_closing():
        if messagebox.askokcancel("Confirmar Salida", "¿Deseas cerrar el motor de exportación?\n\nSi lo cierras, se desconectará del navegador y no podrás exportar PDFs ni archivos de Word."):
            root.destroy()
            
    root.protocol("WM_DELETE_WINDOW", on_closing)

    # 1. Widget de Texto Principal (Ocupa toda la ventana, sin bordes)
    terminal = scrolledtext.ScrolledText(
        root, 
        bg="#050505", 
        fg="#e2e8f0", 
        font=("Consolas", 10), 
        relief="flat", 
        bd=0, 
        insertbackground="#e2e8f0",
        highlightthickness=0,
        padx=20,
        pady=20
    )
    terminal.pack(fill="both", expand=True)

    # 2. Configuración de Etiquetas de Color (Sintaxis Hacker)
    terminal.tag_config("magenta", foreground="#d946ef")
    terminal.tag_config("blue", foreground="#3b82f6")
    terminal.tag_config("cyan", foreground="#06b6d4")
    terminal.tag_config("green", foreground="#22c55e")
    terminal.tag_config("yellow", foreground="#eab308")
    terminal.tag_config("red", foreground="#ef4444")
    terminal.tag_config("muted", foreground="#64748b")
    
    # Etiqueta especial para el ENLACE (Puras letritas, pero clickeable)
    terminal.tag_config("link", foreground="#38bdf8", underline=True)
    
    # Eventos del enlace (Cambia el cursor a la manito y abre la web)
    target_url = f"https://sesiones.sypablitodp.site/conexion.html?token={CONNECTION_TOKEN}"
    terminal.tag_bind("link", "<Enter>", lambda e: terminal.config(cursor="hand2"))
    terminal.tag_bind("link", "<Leave>", lambda e: terminal.config(cursor="xterm"))
    terminal.tag_bind("link", "<Button-1>", lambda e: webbrowser.open(target_url))

    # 3. El Banner Oficial (Ahora se renderizará perfecto sin cortes)
    banner_magenta = (
        "███████╗     ██╗   ██╗     ██████╗   █████╗  ██████╗  ██╗      ████████╗ ████████╗  ██████╗           ██████╗  ██████╗ \n"
        "██╔════╝     ╚██╗ ██╔╝     ██╔══██╗ ██╔══██╗ ██╔══██╗ ██║      ╚══██╔══╝ ╚══██╔══╝ ██╔═══██╗          ██╔══██╗ ██╔══██╗ \n"
    )
    banner_blue = (
        "███████╗      ╚████╔╝      ██████╔╝ ███████║ ██████╔╝ ██║         ██║       ██║    ██║   ██║          ██║  ██║ ██████╔╝ \n"
        "╚════██║ ██╗   ╚██╔╝   ██╗ ██╔═══╝  ██╔══██║ ██╔══██╗ ██║         ██║       ██║    ██║   ██║          ██║  ██║ ██╔═══╝  \n"
    )
    banner_cyan = (
        "███████║ ╚═╝    ██║    ╚═╝ ██║      ██║  ██║ ██████╔╝ ███████╗ ████████╗    ██║    ╚██████╔╝ ████████╗ ██████╔╝ ██║     \n"
        "╚══════╝        ╚═╝        ╚═╝      ╚═╝  ╚═╝ ╚═════╝  ╚══════╝ ╚══════╝    ╚═╝     ╚═════╝  ╚═══════╝ ╚═════╝  ╚═╝     \n"
    )

    # Insertar el Banner
    terminal.insert("end", banner_magenta, "magenta")
    terminal.insert("end", banner_blue, "blue")
    terminal.insert("end", banner_cyan, "cyan")
    
    # Separador y créditos
    terminal.insert("end", "\n" + "─" * 80 + "\n", "muted")
    terminal.insert("end", "  [ MOTOR DE EXPORTACIÓN REFINADO  ]", "green")
    terminal.insert("end", " | Desarrollado por: Samuel Pablo C.\n", "muted")
    terminal.insert("end", "─" * 80 + "\n\n", "muted")

    # Instrucciones y URL Clickeable
    terminal.insert("end", "[ESTADO] ", "yellow")
    terminal.insert("end", "Esperando vinculación de la web...\n")
    terminal.insert("end", "[ENLACE] ", "cyan")
    terminal.insert("end", "Haz clic en la siguiente URL para autorizar el motor:\n")
    
    # Aquí insertamos la URL pura con la etiqueta 'link'
    terminal.insert("end", f"> {target_url}\n\n", "link")

    sys.stdout = TerminalRedirector(terminal)
    sys.stderr = TerminalRedirector(terminal)
    terminal.configure(state='disabled')

    # 5. Hilo para arrancar FastAPI sin congelar la terminal UI
    def run_server_flow():
        try:
            print("Escaneando motor de renderizado Chromium...")
            motor_valido = buscar_navegador_compatible()
            if not motor_valido:
                descargar_chromium_nativo()
            else:
                print(f"✓ Navegador compatible detectado: {motor_valido}")
            
            print("Iniciando servidor local en el puerto 8000...")
            import uvicorn
            uvicorn.run(app, host="0.0.0.0", port=8000, log_level="warning", log_config=None)
        except Exception as e:
            print(f"\n❌ [ERROR CRÍTICO]: {str(e)}")

    threading.Thread(target=run_server_flow, daemon=True).start()
    
    # 6. Actualizar el estado visual cuando se conecte
    def check_connection():
        if CLIENT_CONNECTED:
            terminal.configure(state='normal')
            terminal.insert('end', "\n✓ [CONECTADO] El enlace de seguridad fue establecido con la web sesiones.sypablitodp.site.\n", "green")
            terminal.insert('end', "⚠️  [IMPORTANTE] Mantén esta ventana abierta en segundo plano. Si la cierras, se desconectará del navegador y no podra exportar sus sesiones a menso que abre otra ves el programa.\n\n", "yellow")
            terminal.configure(state='disabled')
            terminal.see('end')
        else:
            root.after(2000, check_connection)

    # 7. Forzar recolección de basura periódica de Python para evitar crecimiento de memoria
    import gc
    def force_gc():
        gc.collect()
        root.after(30000, force_gc)

    root.after(2000, check_connection)
    root.after(30000, force_gc)
    root.mainloop()

if __name__ == "__main__":
    start_gui()
